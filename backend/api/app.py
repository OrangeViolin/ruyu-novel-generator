from fastapi import FastAPI, Depends, HTTPException, BackgroundTasks, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from datetime import datetime
import concurrent.futures
from sqlalchemy import func
from contextlib import contextmanager
import sys
import os
import re
import json

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from backend.database.models import (
    init_db, get_db, ExampleAnalysis, NovelProject, PlotModule, CrawlTask, Submission,
    Character, PlotOutline, ChapterDraft,
    Agent, AgentExecution, AgentVersion, AgentShare, ReferenceMaterial, WritingStyle,
    ChannelAgent, Manuscript, ManuscriptStep, LongNovelMapping, ImitationProject, ImitationStep
)
from backend.ai.ai_factory import AIClientFactory
from backend.generator.novel_builder import NovelBuilder
from backend.generator.plot_assembler import PlotAssembler
from backend.crawler.zhihu_crawler import ZhihuCrawler
from backend.crawler.xiaohongshu_crawler import XiaohongshuCrawler
from backend.crawler.jinjiang_crawler import JinjiangCrawler
from backend.crawler.qidian_crawler import QidianCrawler
from backend.crawler.feilu_crawler import FeiluCrawler
from backend.crawler.k17_crawler import K17Crawler
from backend.analyzer.plot_extractor import PlotExtractor
from backend.analyzer.emotion_analyzer import EmotionAnalyzer
from backend.generator.expansion_engine import ExpansionEngine
from config.settings import settings

# 初始化数据库
init_db()

# 创建FastAPI应用（禁用自动文档）
app = FastAPI(title="网文生成工具", version="1.0.0", docs_url=None, redoc_url=None)

# 配置CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 添加请求日志中间件
@app.middleware("http")
async def log_requests(request, call_next):
    import time
    start_time = time.time()

    # 读取请求体（如果有的话）
    body = None
    if request.method in ["POST", "PUT", "PATCH"]:
        try:
            body = await request.body()
            print(f"\n{'='*50}")
            print(f"{request.method} {request.url.path}")
            print(f"Content-Length: {len(body) if body else 0} bytes")
            if body and len(body) < 10000:
                print(f"Body preview: {body[:500]}")
            print(f"{'='*50}\n")
        except Exception as e:
            print(f"Error reading body: {e}")

    response = await call_next(request)

    process_time = time.time() - start_time
    print(f"{request.method} {request.url.path} - Status: {response.status_code} - Time: {process_time:.2f}s")

    return response

# 挂载静态文件
from starlette.responses import Response
from starlette.staticfiles import StaticFiles

class CachedStaticFiles(StaticFiles):
    async def get_response(self, path: str, scope):
        response = await super().get_response(path, scope)
        if isinstance(response, Response):
            # 禁用浏览器缓存
            response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
        return response

app.mount("/static", CachedStaticFiles(directory="frontend/static"), name="static")

# 初始化核心组件
ai_client = AIClientFactory.get_client()
plot_assembler = PlotAssembler()
novel_builder = NovelBuilder(ai_client=ai_client, plot_assembler=plot_assembler)
plot_extractor = PlotExtractor(ai_client=ai_client)
emotion_analyzer = EmotionAnalyzer()

def get_ai_client(provider: Optional[str] = None, model: Optional[str] = None):
    """根据提供商和模型名称获取AI客户端"""
    return AIClientFactory.get_client(provider or "deepseek", model)


# ========== 数据库辅助函数 ==========

@contextmanager
def get_db_session():
    """数据库会话上下文管理器，确保连接总是被关闭"""
    db_gen = get_db()
    db = next(db_gen)
    try:
        yield db
    finally:
        db.close()


# ========== 数据模型 ==========

class GenerateNovelRequest(BaseModel):
    theme: str
    elements: List[str]
    background: str = "港澳/金牌播报员"
    characters: Optional[Dict[str, Any]] = None
    target_words: int = 10000
    model_provider: Optional[str] = "deepseek"
    model_name: Optional[str] = None


class CrawlRequest(BaseModel):
    source: str  # zhihu, xiaohongshu
    keyword: str
    limit: int = 20


class AnalyzeRequest(BaseModel):
    content: str
    use_ai: bool = True


class PolishRequest(BaseModel):
    content: str
    focus: str = "情绪钩子"
    style: str = "港澳播报员口吻"
    model_provider: Optional[str] = "deepseek"
    model_name: Optional[str] = None


class UpdateChapterRequest(BaseModel):
    project_id: int
    chapter_id: int
    title: Optional[str] = None
    content: Optional[str] = None


class CreateProjectRequest(BaseModel):
    """创建新项目的请求模型 - 用于协作写作系统"""
    name: str  # 项目名称
    theme: Optional[str] = None  # 主题
    background: Optional[str] = None  # 背景设定
    target_words: Optional[int] = 10000  # 目标字数
    genre: Optional[str] = None  # 题材类型
    core_conflict: Optional[str] = None  # 核心冲突
    core_task: Optional[str] = None  # 核心任务


# ========== API路由 ==========

@app.get("/")
async def root():
    """首页 - 返回前端界面"""
    return FileResponse("frontend/index.html")


@app.get("/imitation")
async def imitation_page():
    """仿写页面"""
    return FileResponse("frontend/imitation.html")


@app.get("/api/health")
async def health():
    """健康检查"""
    return {"status": "ok", "timestamp": datetime.now().isoformat()}


# ========== 小说生成相关 ==========

@app.post("/api/novel/generate")
async def generate_novel(request: GenerateNovelRequest, background_tasks: BackgroundTasks):
    """生成小说"""
    try:
        # 创建项目
        db = next(get_db())
        project = NovelProject(
            name=request.theme,
            status="generating",
            outline={"theme": request.theme, "elements": request.elements},
            characters=request.characters or {},
            chapters=[]
        )
        db.add(project)
        db.commit()
        db.refresh(project)

        # 后台生成
        project_id = project.id  # 保存项目ID
        model_provider = request.model_provider
        model_name = request.model_name
        
        def generate_task():
            # 在后台任务中创建新的数据库会话
            task_db = next(get_db())
            try:
                print(f"[生成任务] 项目 {project_id} 开始生成...")
                
                # 获取指定的AI客户端
                current_ai_client = get_ai_client(model_provider, model_name)
                # 更新builder的客户端
                novel_builder.ai_client = current_ai_client

                novel = novel_builder.build_novel(
                    theme=request.theme,
                    elements=request.elements,
                    characters=request.characters,
                    background=request.background
                )

                print(f"[生成任务] 项目 {project_id} 生成完成")

                # 重新查询项目对象
                task_project = task_db.query(NovelProject).filter(NovelProject.id == project_id).first()
                if task_project:
                    # 更新项目
                    task_project.status = "completed"
                    task_project.outline = novel.get("outline", {})
                    task_project.characters = novel.get("characters", {})
                    task_project.chapters = novel.get("chapters", [])
                    task_project.word_count = novel.get("total_words", 0)
                    task_project.updated_at = datetime.now()
                    
                    # 保存到稿件表 (Manuscript)
                    try:
                        manuscript = Manuscript(
                            project_id=project_id,
                            title=task_project.name,
                            content=task_project.chapters,
                            created_at=datetime.now(),
                            updated_at=datetime.now()
                        )
                        task_db.add(manuscript)
                        print(f"[生成任务] 项目 {project_id} 稿件记录已准备")
                    except Exception as ms_err:
                        print(f"保存稿件记录准备失败: {ms_err}")
                        
                    task_db.commit()
                    print(f"[生成任务] 项目 {project_id} 已更新为完成状态")
                else:
                    print(f"[生成任务] 错误：找不到项目 {project_id}")
            except Exception as e:
                print(f"[生成任务] 项目 {project_id} 生成失败: {e}")
                import traceback
                traceback.print_exc()
                # 更新失败状态
                try:
                    task_project = task_db.query(NovelProject).filter(NovelProject.id == project_id).first()
                    if task_project:
                        task_project.status = "failed"
                        task_db.commit()
                        print(f"[生成任务] 项目 {project_id} 已标记为失败")
                except Exception as update_error:
                    print(f"[生成任务] 更新失败状态时出错: {update_error}")
            finally:
                task_db.close()
                print(f"[生成任务] 项目 {project_id} 任务结束")

        background_tasks.add_task(generate_task)

        return {
            "success": True,
            "project_id": project.id,
            "message": "生成任务已启动，请稍候..."
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/novel/projects")
async def list_projects():
    """获取所有项目"""
    db_gen = get_db()
    db = next(db_gen)
    try:
        projects = db.query(NovelProject).order_by(NovelProject.created_at.desc()).all()
        return {
            "projects": [
                {
                    "id": p.id,
                    "name": p.name,
                    "status": p.status,
                    "type": p.type,
                    "source_manuscript_id": p.source_manuscript_id,
                    "word_count": p.word_count,
                    "created_at": p.created_at.isoformat(),
                    "updated_at": p.updated_at.isoformat()
                }
                for p in projects
            ]
        }
    finally:
        db.close()


@app.post("/api/novel/projects")
async def create_project(request: CreateProjectRequest):
    """创建新项目 - 用于协作写作系统"""
    try:
        db = next(get_db())
        project = NovelProject(
            name=request.name,
            theme=request.theme,
            background=request.background,
            target_words=request.target_words,
            genre=request.genre,
            core_conflict=request.core_conflict,
            core_task=request.core_task,
            status="planning",
            outline={},
            characters={},
            chapters=[],
            word_count=0
        )
        db.add(project)
        db.commit()
        db.refresh(project)

        return {
            "success": True,
            "project_id": project.id,
            "message": "项目创建成功",
            "project": {
                "id": project.id,
                "name": project.name,
                "theme": project.theme,
                "background": project.background,
                "target_words": project.target_words,
                "genre": project.genre,
                "core_conflict": project.core_conflict,
                "core_task": project.core_task,
                "status": project.status,
                "created_at": project.created_at.isoformat()
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/projects/generate-full")
async def generate_full_project():
    """AI生成全套项目设定（仅返回内容，不创建）"""
    try:
        # 生成狗血且独特的设定
        prompt = """请生成一个狗血、独特、出人意料的小说设定。要求：

1. 核心冲突要极其狗血，但要有新意，不要老套路
2. 人物关系要错综复杂，充满反转
3. 秘密要惊天动地
4. 主题可以是：家族秘密、身份错位、复仇、豪门恩怨、医疗事故、校园霸凌反转等

重要：请严格按照以下JSON格式返回，不要添加任何其他文字说明：

```json
{
    "name": "项目名称",
    "theme": "一句话主题",
    "background": "现代都市",
    "genre": "世情文",
    "target_words": 15000,
    "core_conflict": "核心冲突描述200字",
    "core_task": "主角的核心任务",
    "characters": [
        {
            "name": "角色名",
            "role_type": "protagonist",
            "age": 28,
            "gender": "男",
            "personality": "性格特点50字",
            "background": "背景故事100字",
            "motivation": "核心动机50字",
            "secret": "隐藏秘密50字，要出人意料",
            "speech_pattern": "语言风格描述50字",
            "behavior_habits": "行为习惯30字"
        }
    ],
    "outlines": [
        {
            "chapter_number": 1,
            "title": "章节标题",
            "summary": "章节摘要100字",
            "plot_points": ["情节要点1", "情节要点2", "情节要点3"],
            "target_words": 2000,
            "focus_elements": ["情绪钩子", "强冲突"],
            "emotion_arc": "情绪变化50字",
            "characters_involved": ["主角"]
        }
    ]
}
```

狗血创意参考：
- 父亲其实是继母的前男友
- 女儿的亲生父亲是公公
- 主角收养的孩子其实是自己的亲弟弟
- 夫妻两人其实是同父异母的兄妹
- 主角的仇人其实是自己的生父
- 母亲为了保护儿子，嫁给了杀害丈夫的凶手

请生成3个主角/反派和3个章节大纲。语言风格要具体，不要泛泛而谈。"""

        messages = [{"role": "user", "content": prompt}]
        response_content = ai_client._call_api(messages, temperature=0.95)

        # 解析AI返回的JSON
        import json
        import re

        # 尝试直接解析
        project_data = None
        try:
            # 先尝试直接解析整个响应
            project_data = json.loads(response_content.strip())
        except:
            # 如果失败，尝试提取JSON部分
            # 查找第一个 { 和最后一个 }
            start_idx = response_content.find('{')
            end_idx = response_content.rfind('}')

            if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
                json_str = response_content[start_idx:end_idx + 1]
                # 清理可能的markdown代码块标记
                json_str = re.sub(r'```json\s*', '', json_str)
                json_str = re.sub(r'```\s*$', '', json_str)
                # 清理注释
                json_str = re.sub(r'//.*?\n', '\n', json_str)
                json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)

                try:
                    project_data = json.loads(json_str)
                except json.JSONDecodeError as e:
                    # 如果还是失败，尝试修复常见的JSON问题
                    # 比如：末尾多余的逗号
                    json_str = re.sub(r',\s*}', '}', json_str)
                    json_str = re.sub(r',\s*]', ']', json_str)
                    # 移除控制字符
                    json_str = ''.join(char for char in json_str if ord(char) >= 32 or char in '\n\r\t')

                    try:
                        project_data = json.loads(json_str)
                    except:
                        print(f"JSON解析失败: {e}")
                        print(f"提取的JSON字符串: {json_str[:500]}...")
                        raise HTTPException(status_code=500, detail="AI生成的内容格式错误，请重试")

        if not project_data:
            raise HTTPException(status_code=500, detail="AI生成的内容格式错误")

        # 只返回生成的数据，不创建数据库记录
        return {
            "success": True,
            "data": project_data
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/projects/create-from-preview")
async def create_project_from_preview(request: dict):
    """从预览数据创建项目"""
    try:
        project_data = request.get("project")
        if not project_data:
            raise HTTPException(status_code=400, detail="缺少项目数据")

        # 创建项目
        db_gen = get_db()
        db = next(db_gen)
        try:
            project = NovelProject(
                name=project_data["name"],
                theme=project_data["theme"],
                background=project_data.get("background", "现代都市"),
                target_words=project_data.get("target_words", 15000),
                genre=project_data.get("genre", "世情文"),
                core_conflict=project_data["core_conflict"],
                core_task=project_data["core_task"],
                status="planning",
                outline={},
                characters={},
                chapters=[],
                word_count=0
            )
            db.add(project)
            db.commit()
            db.refresh(project)

            # 创建人物
            for char_data in project_data.get("characters", []):
                character = Character(
                    project_id=project.id,
                    name=char_data["name"],
                    role_type=char_data.get("role_type", "supporting"),
                    age=char_data.get("age"),
                    gender=char_data.get("gender"),
                    personality=char_data.get("personality"),
                    background=char_data.get("background"),
                    motivation=char_data.get("motivation"),
                    secret=char_data.get("secret"),
                    speech_pattern=char_data.get("speech_pattern"),
                    behavior_habits=char_data.get("behavior_habits"),
                    source="ai_generated"
                )
                db.add(character)
                db.commit()

            # 创建大纲
            for outline_data in project_data.get("outlines", []):
                outline = PlotOutline(
                    project_id=project.id,
                    level="chapter",
                    chapter_number=outline_data["chapter_number"],
                    title=outline_data["title"],
                    summary=outline_data["summary"],
                    plot_points=outline_data.get("plot_points", []),
                    target_words=outline_data.get("target_words", 2000),
                    focus_elements=outline_data.get("focus_elements", []),
                    emotion_arc=outline_data.get("emotion_arc"),
                    characters_involved=outline_data.get("characters_involved", []),
                    source="ai_generated",
                    status="draft",
                    order=outline_data["chapter_number"]
                )
                db.add(outline)
                db.commit()

            return {
                "success": True,
                "project_id": project.id,
                "message": "项目创建成功！"
            }
        finally:
            db.close()

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/novel/project/{project_id}")
async def get_project(project_id: int):
    """获取项目详情"""
    db = next(get_db())
    project = db.query(NovelProject).filter(NovelProject.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    return {
        "id": project.id,
        "name": project.name,
        "outline": project.outline,
        "characters": project.characters,
        "chapters": project.chapters,
        "status": project.status,
        "word_count": project.word_count,
        "created_at": project.created_at.isoformat(),
        "updated_at": project.updated_at.isoformat()
    }


@app.put("/api/novel/chapter")
async def update_chapter(request: UpdateChapterRequest):
    """更新章节"""
    db = next(get_db())
    project = db.query(NovelProject).filter(NovelProject.id == request.project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    # 更新章节
    for chapter in project.chapters:
        if chapter["id"] == request.chapter_id:
            if request.title:
                chapter["title"] = request.title
            if request.content:
                chapter["content"] = request.content
            chapter["word_count"] = len(chapter.get("content", ""))
            break

    # 更新总字数
    project.word_count = sum(c.get("word_count", 0) for c in project.chapters)
    project.updated_at = datetime.now()
    db.commit()

    return {"success": True, "message": "章节已更新"}


@app.post("/api/novel/polish")
async def polish_text(request: PolishRequest):
    """润色文本"""
    try:
        current_ai_client = get_ai_client(request.model_provider, request.model_name)
        novel_builder.ai_client = current_ai_client
        polished = novel_builder.polish_chapter(
            content=request.content,
            focus=request.focus,
            style=request.style
        )
        return {"success": True, "polished": polished}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/novel/export/{project_id}")
async def export_novel_to_word(project_id: int):
    """导出小说为Word文档"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from fastapi.responses import Response

        db = next(get_db())
        project = db.query(NovelProject).filter(NovelProject.id == project_id).first()

        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        # 创建Word文档
        doc = Document()

        # 设置文档标题
        title = doc.add_heading(project.name, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加元信息
        info_para = doc.add_paragraph()
        info_para.add_run(f"字数：{project.word_count} 字\n").bold = True
        info_para.add_run(f"创建时间：{project.created_at.strftime('%Y年%m月%d日')}\n")
        info_para.add_run(f"更新时间：{project.updated_at.strftime('%Y年%m月%d日')}")

        # 添加分割线
        doc.add_paragraph('_' * 50)

        # 添加章节
        for chapter in project.chapters:
            # 章节标题
            doc.add_heading(chapter.get('title', '未命名章节'), 1)

            # 章节内容
            content = chapter.get('content', '')

            # 将内容分段落
            paragraphs = content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    run = para.add_run(para_text.strip())
                    # 设置字体
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
                    # 设置段落首行缩进
                    para.paragraph_format.first_line_indent = Inches(0.3)
                    para.paragraph_format.line_spacing = 1.5

            # 章节之间添加空行
            doc.add_paragraph()

        # 保存到内存
        from io import BytesIO
        from urllib.parse import quote
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # 返回Word文件 (使用URL编码解决中文文件名问题)
        encoded_filename = quote(f"{project.name}.docx")
        return Response(
            content=buffer.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename*=UTF-8\'\'{encoded_filename}'
            }
        )

    except ImportError:
        raise HTTPException(status_code=500, detail="请先安装python-docx库: pip install python-docx")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 例文拆解相关 ==========

class FetchContentRequest(BaseModel):
    url: str


@app.post("/api/fetch-content")
async def fetch_content_from_url(request: FetchContentRequest):
    """从URL获取内容"""
    try:
        import requests
        from bs4 import BeautifulSoup
        import re

        # 发送请求获取页面
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }

        response = requests.get(request.url, headers=headers, timeout=15)
        response.raise_for_status()
        response.encoding = response.apparent_encoding

        soup = BeautifulSoup(response.text, 'html.parser')

        # 提取标题
        title = ""
        # 优先从meta标签获取标题
        og_title = soup.find('meta', property='og:title')
        if og_title:
            title = og_title.get('content', '')
        else:
            title_tag = soup.find('h1')
            if not title_tag:
                title_tag = soup.find('title')
            if title_tag:
                title = title_tag.get_text(strip=True)
                # 清理标题中的后缀
                title = title.split('_')[0].split('|')[0].split('-')[0].strip()

        # 提取正文内容
        content = ""

        # 针对网易文章的特定处理
        if '163.com' in request.url or 'dy/article' in request.url:
            # 网易文章的特定选择器
            # 首先尝试查找article标签
            article = soup.find('article')

            # 如果找不到article，尝试查找包含"post"的div
            if not article:
                article = soup.find('div', class_=lambda x: x and any('post' in str(c).lower() for c in x))

            # 如果还是找不到，尝试查找主体内容div
            if not article:
                for selector in ['.post_text', '.post-text-b', '.text', '.content', '.article-content']:
                    article = soup.select_one(selector)
                    if article:
                        break

            if article:
                # 移除article内的无用元素
                for tag in article(['script', 'style', 'iframe', 'noscript', 'nav', 'header', 'footer', 'aside']):
                    tag.decompose()

                # 移除广告和导航div
                for element in article.find_all('div'):
                    class_attr = ' '.join(element.get('class', []))
                    id_attr = element.get('id', '')
                    if any(nav_word in class_attr.lower() or nav_word in id_attr.lower()
                           for nav_word in ['nav', 'menu', 'ad', 'banner', 'toolbar', 'footer', 'ggw', 'ad-', 'guanggao']):
                        element.decompose()

                # 提取所有段落
                paragraphs = []
                for p in article.find_all('p'):
                    text = p.get_text(strip=True)
                    if text and len(text) > 5:
                        # 过滤掉明显的导航和广告文本
                        skip_keywords = [
                            '网易首页', '快速导航', '返回网易首页', '下载网易新闻客户端',
                            '特别声明：以上内容', 'Notice: The content above',
                            '阅读下一篇', '相关推荐', '热点推荐',
                            '分享至好友和朋友圈',
                            '用微信扫码', '举报',
                            '来源:', '跟贴', '转载请注明', '本文由', '关注公众号',
                            '长按识别二维码', '点击查看更多',
                            '查看网易地图', '登录', '注册免费邮箱', '申请入驻',
                            '###', '####', '#####',  # Markdown标题标记
                            '一梦春风拂柳近', '清禅幽送芳菲来'  # 作者信息
                        ]

                        should_skip = False
                        for keyword in skip_keywords:
                            if keyword in text:
                                should_skip = True
                                break

                        # 跳过单独的导航类别
                        if text.strip() in ['新闻', '体育', '娱乐', '财经', '汽车', '科技', '时尚', '房产', '教育', '手机', '数码']:
                            should_skip = True

                        # 跳过时间戳行
                        if re.match(r'^\d{4}-\d{2}-\d{2}.*来源:', text):
                            should_skip = True

                        # 检查是否包含过多#符号（可能是导航标题）
                        if text.count('#') >= 3:
                            should_skip = True

                        # 检查是否包含过多短横线分隔的导航项
                        if text.count(' - ') > 2 and len(text) < 100:
                            should_skip = True

                        if not should_skip:
                            paragraphs.append(text)

                # 如果提取到了足够的段落，使用这些段落
                if paragraphs and len('\n\n'.join(paragraphs)) > 200:
                    content = '\n\n'.join(paragraphs)

        # 如果网易特定方法失败，使用通用方法
        if not content or len(content) < 100:
            # 通用回退方法：获取所有p标签
            # 先移除明显的无用标签
            for tag in soup(['script', 'style', 'nav', 'header', 'footer', 'aside', 'iframe', 'noscript']):
                tag.decompose()

            # 获取所有p标签
            all_p = soup.find_all('p')
            paragraphs = []

            for p in all_p:
                text = p.get_text(strip=True)
                if text and len(text) > 10:
                    # 过滤导航和广告
                    skip_keywords = [
                        '网易首页', '快速导航', '返回', '下载', '特别声明',
                        '相关推荐', '热点推荐', '分享至', '举报',
                        '###', '####', '来源:', '跟贴'
                    ]
                    if not any(kw in text for kw in skip_keywords):
                        # 跳过单独的导航类别
                        if text.strip() not in ['新闻', '体育', '娱乐', '财经', '汽车', '科技', '时尚', '房产', '教育', '手机', '数码']:
                            paragraphs.append(text)

            if paragraphs:
                # 取所有段落，让清理函数来过滤
                content = '\n\n'.join(paragraphs)

        # 最后的回退：如果还是没有内容，尝试直接获取body文本
        if not content or len(content) < 50:
            body = soup.find('body')
            if body:
                # 移除所有无用元素
                for tag in body(['script', 'style', 'nav', 'header', 'footer', 'aside', 'iframe', 'noscript']):
                    tag.decompose()

                content = body.get_text(separator='\n', strip=True)

        # 清理内容
        content = _clean_content(content)

        word_count = len(content)

        if word_count < 50:
            return {
                "success": False,
                "error": f"提取内容过少（{word_count}字），该网站可能不支持自动提取"
            }

        return {
            "success": True,
            "title": title,
            "content": content,
            "author": "",
            "word_count": word_count
        }

    except requests.RequestException as e:
        return {
            "success": False,
            "error": f"无法获取页面: {str(e)}"
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"提取失败: {str(e)}"
        }


def _extract_text_from_element(element):
    """从元素中提取文本"""
    texts = []
    for child in element.descendants:
        if child.name == 'p':
            text = child.get_text(strip=True)
            if text:
                texts.append(text)
        elif child.name == 'br':
            texts.append("")

    return '\n\n'.join(texts) if texts else element.get_text(separator='\n', strip=True)


def _clean_content(text):
    """清理内容"""
    if not text:
        return ""

    # 移除常见的无用文本（使用更精确的模式）
    noise_patterns = [
        r'特别声明：以上内容[\s\S]*?为自媒体平台.*?仅提供信息存储服务',
        r'Notice: The content above[\s\S]*?only provides information storage services',
        r'阅读下一篇',
        r'### \s*\n',  # 单独的 ### 标题行
        r'声明.*?版权归.*',
        r'本文由.*原创',
        r'转载请注明.*',
        r'长按识别二维码.*',
        r'关注公众号.*',
        r'分享至好友和朋友圈',
        r'用微信扫码.*',
        r'举报\d*$',
        r'查看网易地图',
        r'登录\s*\n注册免费邮箱',
    ]

    for pattern in noise_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.DOTALL)

    # 移除空行
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

    # 分行并智能清理
    lines = text.split('\n')
    clean_lines = []

    # 定义导航关键词
    nav_keywords = [
        '网易首页', '应用', '快速导航', '返回网易首页', '下载网易新闻客户端',
        '相关推荐', '热点推荐', '特别声明', '来源:', '跟贴'
    ]

    # 定义单独的导航类别（如"新闻"、"体育"等后跟大量选项）
    nav_categories = ['新闻', '体育', '娱乐', '财经', '汽车', '科技', '时尚', '房产', '教育', '手机', '数码']

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 跳过包含明显导航关键词的行
        if any(keyword in line for keyword in nav_keywords):
            continue

        # 跳过单独的导航类别行
        if line in nav_categories:
            continue

        # 跳过纯数字或很短的行（但保留中文数字或混合内容）
        if len(line) <= 3 and line.isdigit():
            continue

        # 跳过时间戳行（如：2025-08-31 06:30:44 来源: xxx）
        if re.match(r'^\d{4}-\d{2}-\d{2}.*来源:', line):
            continue

        # 跳过包含多个短横线分隔的导航项（如：新闻 - 体育 - 娱乐）
        if line.count('-') > 3 and len(line) < 100:
            continue

        # 保留合理的行（超过5个字符的文本）
        if len(line) > 5:
            clean_lines.append(line)

    # 保留更多内容，不仅仅是前50行
    return '\n\n'.join(clean_lines).strip()



class ExampleAnalysisRequest(BaseModel):
    title: str
    source_url: Optional[str] = None
    content: str
    analysis_title: Optional[str] = None
    core_conflict: Optional[str] = None
    information_gap: Optional[str] = None
    core_task: Optional[str] = None
    character_profile: Optional[str] = None
    notes: Optional[str] = None
    tags: Optional[List[str]] = None


@app.get("/api/examples")
async def list_examples():
    """获取所有例文拆解笔记"""
    db = next(get_db())
    examples = db.query(ExampleAnalysis).order_by(ExampleAnalysis.updated_at.desc()).all()

    return {
        "examples": [
            {
                "id": e.id,
                "title": e.title,
                "source_url": e.source_url,
                "content_preview": e.content[:200] + "..." if e.content and len(e.content) > 200 else e.content,
                "tags": e.tags or [],
                "created_at": e.created_at.isoformat(),
                "updated_at": e.updated_at.isoformat()
            }
            for e in examples
        ]
    }


@app.get("/api/examples/{example_id}")
async def get_example(example_id: int):
    """获取单个例文拆解详情"""
    db = next(get_db())
    example = db.query(ExampleAnalysis).filter(ExampleAnalysis.id == example_id).first()

    if not example:
        raise HTTPException(status_code=404, detail="例文不存在")

    return {
        "id": example.id,
        "title": example.title,
        "source_url": example.source_url,
        "content": example.content,
        "analysis_title": example.analysis_title,
        "core_conflict": example.core_conflict,
        "information_gap": example.information_gap,
        "core_task": example.core_task,
        "character_profile": example.character_profile,
        "notes": example.notes,
        "tags": example.tags or [],
        "created_at": example.created_at.isoformat(),
        "updated_at": example.updated_at.isoformat()
    }


@app.post("/api/examples")
async def create_example(request: ExampleAnalysisRequest):
    """创建新的例文拆解笔记"""
    try:
        db = next(get_db())

        example = ExampleAnalysis(
            title=request.title,
            source_url=request.source_url,
            content=request.content,
            analysis_title=request.analysis_title,
            core_conflict=request.core_conflict,
            information_gap=request.information_gap,
            core_task=request.core_task,
            character_profile=request.character_profile,
            notes=request.notes,
            tags=request.tags or []
        )

        db.add(example)
        db.commit()
        db.refresh(example)

        return {"success": True, "id": example.id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/examples/{example_id}")
async def update_example(example_id: int, request: ExampleAnalysisRequest):
    """更新例文拆解笔记"""
    try:
        db = next(get_db())
        example = db.query(ExampleAnalysis).filter(ExampleAnalysis.id == example_id).first()

        if not example:
            raise HTTPException(status_code=404, detail="例文不存在")

        example.title = request.title
        example.source_url = request.source_url
        example.content = request.content
        example.analysis_title = request.analysis_title
        example.core_conflict = request.core_conflict
        example.information_gap = request.information_gap
        example.core_task = request.core_task
        example.character_profile = request.character_profile
        example.notes = request.notes
        example.tags = request.tags or []
        example.updated_at = datetime.now()

        db.commit()

        return {"success": True, "message": "例文已更新"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/examples/{example_id}")
async def delete_example(example_id: int):
    """删除例文拆解笔记"""
    try:
        db = next(get_db())
        example = db.query(ExampleAnalysis).filter(ExampleAnalysis.id == example_id).first()

        if not example:
            raise HTTPException(status_code=404, detail="例文不存在")

        db.delete(example)
        db.commit()

        return {"success": True, "message": "例文已删除"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 分析相关 ==========

@app.post("/api/analyze/plot")
async def analyze_plot(request: AnalyzeRequest):
    """分析情节元素"""
    try:
        if request.use_ai:
            result = plot_extractor.extract_by_ai(request.content)
        else:
            result = plot_extractor.extract_by_rules(request.content)

        return {"success": True, "result": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/analyze/emotion")
async def analyze_emotion(request: AnalyzeRequest):
    """分析情绪"""
    try:
        result = emotion_analyzer.analyze(request.content)
        hooks = emotion_analyzer.detect_emotional_hooks(request.content)

        return {
            "success": True,
            "emotion": result,
            "hooks": hooks
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 投稿管理相关 ==========

class SubmissionRequest(BaseModel):
    project_id: int
    platform: str
    pen_name: str
    book_name: str
    theme: str
    status: str = "pending"
    submission_date: Optional[str] = None
    notes: Optional[str] = None
    generated_title: Optional[str] = None
    generated_intro: Optional[str] = None


@app.get("/api/submissions")
async def list_submissions():
    """获取所有投稿记录"""
    db = next(get_db())
    submissions = db.query(Submission).order_by(Submission.created_at.desc()).all()

    results = []
    for sub in submissions:
        # 获取项目字数
        project = db.query(NovelProject).filter(NovelProject.id == sub.project_id).first()
        word_count = project.word_count if project else 0

        results.append({
            "id": sub.id,
            "project_id": sub.project_id,
            "platform": sub.platform,
            "pen_name": sub.pen_name,
            "book_name": sub.book_name,
            "theme": sub.theme,
            "word_count": word_count,
            "status": sub.status,
            "submission_date": sub.submission_date,
            "notes": sub.notes,
            "generated_title": sub.generated_title,
            "generated_intro": sub.generated_intro,
            "created_at": sub.created_at.isoformat(),
            "updated_at": sub.updated_at.isoformat()
        })

    return {"submissions": results}


@app.post("/api/submissions")
async def create_submission(request: SubmissionRequest):
    """创建投稿记录"""
    try:
        db = next(get_db())

        # 获取项目字数
        project = db.query(NovelProject).filter(NovelProject.id == request.project_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        submission = Submission(
            project_id=request.project_id,
            platform=request.platform,
            pen_name=request.pen_name,
            book_name=request.book_name,
            theme=request.theme,
            word_count=project.word_count,
            status=request.status,
            submission_date=request.submission_date,
            notes=request.notes,
            generated_title=request.generated_title,
            generated_intro=request.generated_intro
        )

        db.add(submission)
        db.commit()
        db.refresh(submission)

        return {"success": True, "id": submission.id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/submissions/{submission_id}")
async def get_submission(submission_id: int):
    """获取投稿详情"""
    db = next(get_db())
    submission = db.query(Submission).filter(Submission.id == submission_id).first()

    if not submission:
        raise HTTPException(status_code=404, detail="投稿记录不存在")

    # 获取项目字数
    project = db.query(NovelProject).filter(NovelProject.id == submission.project_id).first()
    word_count = project.word_count if project else 0

    return {
        "id": submission.id,
        "project_id": submission.project_id,
        "platform": submission.platform,
        "pen_name": submission.pen_name,
        "book_name": submission.book_name,
        "theme": submission.theme,
        "word_count": word_count,
        "status": submission.status,
        "submission_date": submission.submission_date,
        "notes": submission.notes,
        "generated_title": submission.generated_title,
        "generated_intro": submission.generated_intro,
        "created_at": submission.created_at.isoformat(),
        "updated_at": submission.updated_at.isoformat()
    }


@app.put("/api/submissions/{submission_id}")
async def update_submission(submission_id: int, request: SubmissionRequest):
    """更新投稿记录"""
    try:
        db = next(get_db())
        submission = db.query(Submission).filter(Submission.id == submission_id).first()

        if not submission:
            raise HTTPException(status_code=404, detail="投稿记录不存在")

        # 获取项目字数
        project = db.query(NovelProject).filter(NovelProject.id == request.project_id).first()
        word_count = project.word_count if project else 0

        submission.project_id = request.project_id
        submission.platform = request.platform
        submission.pen_name = request.pen_name
        submission.book_name = request.book_name
        submission.theme = request.theme
        submission.word_count = word_count
        submission.status = request.status
        submission.submission_date = request.submission_date
        submission.notes = request.notes
        submission.generated_title = request.generated_title
        submission.generated_intro = request.generated_intro
        submission.updated_at = datetime.now()

        db.commit()

        return {"success": True, "message": "投稿已更新"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/submissions/{submission_id}")
async def delete_submission(submission_id: int):
    """删除投稿记录"""
    try:
        db = next(get_db())
        submission = db.query(Submission).filter(Submission.id == submission_id).first()

        if not submission:
            raise HTTPException(status_code=404, detail="投稿记录不存在")

        db.delete(submission)
        db.commit()

        return {"success": True, "message": "投稿已删除"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 人机协作 - 人物管理 ==========

class CharacterRequest(BaseModel):
    project_id: int
    name: str
    role_type: str = "supporting"  # protagonist, antagonist, supporting
    age: Optional[int] = None
    gender: Optional[str] = None
    appearance: Optional[str] = None
    personality: Optional[str] = None
    background: Optional[str] = None
    motivation: Optional[str] = None
    secret: Optional[str] = None
    relationships: Optional[List[Dict]] = None
    speech_pattern: Optional[str] = None
    behavior_habits: Optional[str] = None
    emotional_triggers: Optional[str] = None
    notes: Optional[str] = None


@app.get("/api/projects/{project_id}/characters")
async def list_characters(project_id: int):
    """获取项目的所有人物"""
    db = next(get_db())
    characters = db.query(Character).filter(
        Character.project_id == project_id
    ).order_by(Character.id).all()

    return {
        "characters": [
            {
                "id": c.id,
                "name": c.name,
                "role_type": c.role_type,
                "age": c.age,
                "gender": c.gender,
                "appearance": c.appearance,
                "personality": c.personality,
                "background": c.background,
                "motivation": c.motivation,
                "secret": c.secret,
                "relationships": c.relationships or [],
                "speech_pattern": c.speech_pattern,
                "source": c.source,
                "notes": c.notes
            }
            for c in characters
        ]
    }


@app.post("/api/characters")
async def create_character(request: CharacterRequest):
    """创建人物"""
    try:
        db = next(get_db())

        character = Character(
            project_id=request.project_id,
            name=request.name,
            role_type=request.role_type,
            age=request.age,
            gender=request.gender,
            appearance=request.appearance,
            personality=request.personality,
            background=request.background,
            motivation=request.motivation,
            secret=request.secret,
            relationships=request.relationships,
            speech_pattern=request.speech_pattern,
            behavior_habits=request.behavior_habits,
            emotional_triggers=request.emotional_triggers,
            source="manual",
            notes=request.notes
        )

        db.add(character)
        db.commit()
        db.refresh(character)

        return {"success": True, "id": character.id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class GenerateCharacterRequest(BaseModel):
    project_id: int
    role_type: str  # protagonist, antagonist, supporting
    theme: str  # 主题/题材
    elements: List[str]  # 必须包含的元素
    reference: Optional[str] = None  # 参考描述


@app.post("/api/characters/generate")
async def generate_character(request: GenerateCharacterRequest):
    """AI辅助生成人物"""
    try:
        db = next(get_db())

        # 调用AI生成人物设定
        prompt = f"""请根据以下信息，生成一个详细的{request.role_type}角色设定：

主题：{request.theme}
必须包含的元素：{', '.join(request.elements)}
角色类型：{request.role_type}
参考描述：{request.reference or '无'}

请生成以下内容（JSON格式）：
{{
    "name": "角色姓名",
    "age": 年龄,
    "gender": "性别",
    "appearance": "外貌描述（100字左右）",
    "personality": "性格特点（200字左右）",
    "background": "背景故事（300字左右）",
    "motivation": "核心动机/目标",
    "secret": "隐藏的秘密（用于后续反转）",
    "speech_pattern": "语言风格描述（用于AI生成对话）",
    "behavior_habits": "行为习惯特点",
    "emotional_triggers": "情绪触发点"
}}

要求：
1. 人物要有鲜明特点，避免刻板印象
2. 秘密要能推动剧情发展
3. 语言风格要具体，便于AI模仿
"""

        messages = [{"role": "user", "content": prompt}]
        response = ai_client._call_api(messages, temperature=0.8)

        # 解析AI返回的JSON
        import json
        try:
            # 尝试提取JSON部分
            start_idx = response.find('{')
            end_idx = response.rfind('}') + 1
            if start_idx >= 0 and end_idx > start_idx:
                json_str = response[start_idx:end_idx]
                character_data = json.loads(json_str)
            else:
                raise ValueError("未找到JSON格式")
        except:
            raise HTTPException(status_code=500, detail="AI返回格式错误")

        # 创建人物记录
        character = Character(
            project_id=request.project_id,
            name=character_data.get("name", "未命名"),
            role_type=request.role_type,
            age=character_data.get("age"),
            gender=character_data.get("gender"),
            appearance=character_data.get("appearance"),
            personality=character_data.get("personality"),
            background=character_data.get("background"),
            motivation=character_data.get("motivation"),
            secret=character_data.get("secret"),
            speech_pattern=character_data.get("speech_pattern"),
            behavior_habits=character_data.get("behavior_habits"),
            emotional_triggers=character_data.get("emotional_triggers"),
            source="ai_generated"
        )

        db.add(character)
        db.commit()
        db.refresh(character)

        return {
            "success": True,
            "id": character.id,
            "character": {
                "id": character.id,
                "name": character.name,
                "role_type": character.role_type,
                **character_data
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/characters/{character_id}")
async def update_character(character_id: int, request: CharacterRequest):
    """更新人物"""
    try:
        db = next(get_db())
        character = db.query(Character).filter(Character.id == character_id).first()

        if not character:
            raise HTTPException(status_code=404, detail="人物不存在")

        character.name = request.name
        character.role_type = request.role_type
        character.age = request.age
        character.gender = request.gender
        character.appearance = request.appearance
        character.personality = request.personality
        character.background = request.background
        character.motivation = request.motivation
        character.secret = request.secret
        character.relationships = request.relationships
        character.speech_pattern = request.speech_pattern
        character.behavior_habits = request.behavior_habits
        character.emotional_triggers = request.emotional_triggers
        character.notes = request.notes
        character.updated_at = datetime.now()

        db.commit()

        return {"success": True, "message": "人物已更新"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/characters/{character_id}")
async def delete_character(character_id: int):
    """删除人物"""
    try:
        db = next(get_db())
        character = db.query(Character).filter(Character.id == character_id).first()

        if not character:
            raise HTTPException(status_code=404, detail="人物不存在")

        db.delete(character)
        db.commit()

        return {"success": True, "message": "人物已删除"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 人机协作 - 情节大纲管理 ==========

class PlotOutlineRequest(BaseModel):
    project_id: int
    level: str  # story, chapter, scene
    parent_id: Optional[int] = None
    chapter_number: Optional[int] = None
    title: str
    summary: str
    plot_points: Optional[List[str]] = None
    target_words: Optional[int] = None
    focus_elements: Optional[List[str]] = None
    emotion_arc: Optional[str] = None
    characters_involved: Optional[List[str]] = None
    notes: Optional[str] = None


@app.get("/api/projects/{project_id}/outlines")
async def list_plot_outlines(project_id: int, level: Optional[str] = None):
    """获取项目的情节大纲"""
    db = next(get_db())
    query = db.query(PlotOutline).filter(PlotOutline.project_id == project_id)

    if level:
        query = query.filter(PlotOutline.level == level)

    outlines = query.order_by(PlotOutline.order).all()

    return {
        "outlines": [
            {
                "id": o.id,
                "level": o.level,
                "parent_id": o.parent_id,
                "chapter_number": o.chapter_number,
                "title": o.title,
                "summary": o.summary,
                "plot_points": o.plot_points or [],
                "target_words": o.target_words,
                "focus_elements": o.focus_elements or [],
                "emotion_arc": o.emotion_arc,
                "characters_involved": o.characters_involved or [],
                "status": o.status,
                "source": o.source,
                "notes": o.notes,
                "order": o.order
            }
            for o in outlines
        ]
    }


@app.post("/api/outlines")
async def create_plot_outline(request: PlotOutlineRequest):
    """创建情节大纲"""
    try:
        db = next(get_db())

        # 获取当前最大的order值
        max_order = db.query(func.max(PlotOutline.order)).filter(
            PlotOutline.project_id == request.project_id
        ).scalar() or 0

        outline = PlotOutline(
            project_id=request.project_id,
            level=request.level,
            parent_id=request.parent_id,
            chapter_number=request.chapter_number,
            title=request.title,
            summary=request.summary,
            plot_points=request.plot_points or [],
            target_words=request.target_words,
            focus_elements=request.focus_elements or [],
            emotion_arc=request.emotion_arc,
            characters_involved=request.characters_involved or [],
            source="manual",
            status="draft",
            notes=request.notes,
            order=max_order + 1
        )

        db.add(outline)
        db.commit()
        db.refresh(outline)

        return {"success": True, "id": outline.id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/outlines/{outline_id}")
async def update_plot_outline(outline_id: int, request: PlotOutlineRequest):
    """更新情节大纲"""
    try:
        db = next(get_db())
        outline = db.query(PlotOutline).filter(PlotOutline.id == outline_id).first()

        if not outline:
            raise HTTPException(status_code=404, detail="大纲不存在")

        outline.title = request.title
        outline.summary = request.summary
        outline.plot_points = request.plot_points or []
        outline.target_words = request.target_words
        outline.focus_elements = request.focus_elements or []
        outline.emotion_arc = request.emotion_arc
        outline.characters_involved = request.characters_involved or []
        outline.notes = request.notes
        outline.updated_at = datetime.now()

        db.commit()

        return {"success": True, "message": "大纲已更新"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/outlines/{outline_id}")
async def delete_plot_outline(outline_id: int):
    """删除情节大纲"""
    try:
        db = next(get_db())
        outline = db.query(PlotOutline).filter(PlotOutline.id == outline_id).first()

        if not outline:
            raise HTTPException(status_code=404, detail="大纲不存在")

        db.delete(outline)
        db.commit()

        return {"success": True, "message": "大纲已删除"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 人机协作 - 章节草稿管理 ==========

class GenerateChapterRequest(BaseModel):
    outline_id: int
    temperature: Optional[float] = 0.8
    focus: Optional[str] = "情绪钩子"


@app.post("/api/chapters/generate")
async def generate_chapter_content(request: GenerateChapterRequest):
    """根据大纲生成章节内容"""
    try:
        db = next(get_db())

        # 获取大纲信息
        outline = db.query(PlotOutline).filter(PlotOutline.id == request.outline_id).first()
        if not outline:
            raise HTTPException(status_code=404, detail="大纲不存在")

        # 获取项目信息
        project = db.query(NovelProject).filter(NovelProject.id == outline.project_id).first()
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        # 获取涉及的人物
        characters = db.query(Character).filter(
            Character.project_id == outline.project_id,
            Character.name.in_(outline.characters_involved or [])
        ).all()

        # 构建人物上下文
        character_context = ""
        for char in characters:
            character_context += f"""
【{char.name}】
- 角色：{char.role_type}
- 性格：{char.personality or '未设定'}
- 语言风格：{char.speech_pattern or '未设定'}
- 动机：{char.motivation or '未设定'}
"""

        # 生成章节内容
        prompt = f"""请根据以下大纲生成章节内容：

项目主题：{project.theme}
章节标题：{outline.title}
章节摘要：{outline.summary}
情节要点：{', '.join(outline.plot_points or [])}
重点元素：{', '.join(outline.focus_elements or [])}
情绪弧线：{outline.emotion_arc or '未设定'}

人物设定：
{character_context}

要求：
1. 严格按照大纲展开剧情
2. 重点体现{request.focus or '情绪钩子'}
3. 语言风格要符合人物设定
4. 目标字数约{outline.target_words or 2000}字
5. 保持情节连贯，有情绪张力
6. 对话要符合人物语言风格

请直接输出章节内容，不要有标题、说明等前缀。
"""

        messages = [{"role": "user", "content": prompt}]
        content = ai_client._call_api(messages, temperature=request.temperature)

        # 创建章节草稿
        chapter = ChapterDraft(
            project_id=outline.project_id,
            outline_id=request.outline_id,
            chapter_number=outline.chapter_number,
            title=outline.title,
            content=content,
            word_count=len(content),
            status="draft",
            generation_params={
                "temperature": request.temperature,
                "focus": request.focus
            }
        )

        db.add(chapter)
        db.commit()
        db.refresh(chapter)

        # 更新大纲状态
        outline.status = "generated"
        db.commit()

        return {
            "success": True,
            "chapter_id": chapter.id,
            "content": content,
            "word_count": len(content)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/projects/{project_id}/chapters")
async def list_chapters(project_id: int):
    """获取项目的所有章节"""
    db = next(get_db())
    chapters = db.query(ChapterDraft).filter(
        ChapterDraft.project_id == project_id
    ).order_by(ChapterDraft.chapter_number).all()

    return {
        "chapters": [
            {
                "id": c.id,
                "outline_id": c.outline_id,
                "chapter_number": c.chapter_number,
                "title": c.title,
                "content": c.content,
                "word_count": c.word_count,
                "status": c.status,
                "edit_count": c.edit_count,
                "ai_revision_count": c.ai_revision_count,
                "human_ai_ratio": c.human_ai_ratio,
                "notes": c.notes,
                "created_at": c.created_at.isoformat(),
                "updated_at": c.updated_at.isoformat()
            }
            for c in chapters
        ]
    }


class UpdateChapterContentRequest(BaseModel):
    content: str
    notes: Optional[str] = None


@app.put("/api/chapters/{chapter_id}")
async def update_chapter(chapter_id: int, request: UpdateChapterContentRequest):
    """更新章节内容（人工编辑）"""
    try:
        db = next(get_db())
        chapter = db.query(ChapterDraft).filter(ChapterDraft.id == chapter_id).first()

        if not chapter:
            raise HTTPException(status_code=404, detail="章节不存在")

        old_content = chapter.content
        chapter.content = request.content
        chapter.word_count = len(request.content)

        # 检测是否是人工编辑
        if len(request.content) > len(old_content) * 0.7:  # 如果修改超过30%，认为是人工编辑
            chapter.edit_count = (chapter.edit_count or 0) + 1
            # 简单计算人机比例
            total_edits = (chapter.edit_count or 0) + (chapter.ai_revision_count or 0)
            human_ratio = int((chapter.edit_count or 0) / total_edits * 100) if total_edits > 0 else 50
            chapter.human_ai_ratio = f"{human_ratio}:{100-human_ratio}"

        if request.notes:
            chapter.notes = request.notes

        chapter.updated_at = datetime.now()
        db.commit()

        return {
            "success": True,
            "message": "章节已更新",
            "edit_count": chapter.edit_count,
            "human_ai_ratio": chapter.human_ai_ratio
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class ReviseChapterRequest(BaseModel):
    focus: str = "情绪钩子"
    style: Optional[str] = None
    instructions: Optional[str] = None


@app.post("/api/chapters/{chapter_id}/revise")
async def revise_chapter_with_ai(chapter_id: int, request: ReviseChapterRequest):
    """使用AI润色章节"""
    try:
        db = next(get_db())
        chapter = db.query(ChapterDraft).filter(ChapterDraft.id == chapter_id).first()

        if not chapter:
            raise HTTPException(status_code=404, detail="章节不存在")

        # 获取项目信息
        project = db.query(NovelProject).filter(NovelProject.id == chapter.project_id).first()

        # AI润色
        prompt = f"""请润色以下章节内容：

章节标题：{chapter.title}
主题：{project.theme if project else '未设定'}

原文：
{chapter.content}

润色要求：
1. 重点优化：{request.focus}
2. 保持原有剧情不变
{f"3. 风格要求：{request.style}" if request.style else ""}
{f"4. 额外说明：{request.instructions}" if request.instructions else ""}

请直接输出润色后的内容，不要有说明。
"""

        messages = [{"role": "user", "content": prompt}]
        revised_content = ai_client._call_api(messages, temperature=0.7)

        # 更新章节
        old_content = chapter.content
        chapter.content = revised_content
        chapter.word_count = len(revised_content)
        chapter.ai_revision_count = (chapter.ai_revision_count or 0) + 1
        chapter.status = "revising"
        chapter.updated_at = datetime.now()

        # 计算人机比例（简单估算）
        ai_ratio = min(90, 30 + chapter.ai_revision_count * 10)
        chapter.human_ai_ratio = f"{100-ai_ratio}:{ai_ratio}"

        db.commit()

        return {
            "success": True,
            "content": revised_content,
            "word_count": len(revised_content),
            "ai_revision_count": chapter.ai_revision_count
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 模板相关 ==========

@app.get("/api/templates")
async def get_templates():
    """获取所有情节模板"""
    return {"templates": plot_assembler.templates}


@app.post("/api/templates/fill")
async def fill_template(category: str, subcategory: str, name: str, variables: Dict):
    """填充模板"""
    try:
        template = plot_assembler.get_template(category, subcategory, name)
        if not template:
            raise HTTPException(status_code=404, detail="模板不存在")

        content = plot_assembler.fill_template(template, variables)
        return {"success": True, "content": content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== AI智能体相关 ==========

@app.get("/api/agents")
async def list_agents(category: str = None, agent_type: str = None):
    """获取智能体列表"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            query = db.query(Agent)

            # 筛选条件
            if category:
                query = query.filter(Agent.category == category)
            if agent_type:
                query = query.filter(Agent.agent_type == agent_type)

            # 只返回公开的或用户自己的智能体（这里简化为返回所有）
            agents = query.order_by(Agent.order.asc(), Agent.created_at.desc()).all()

            return {
                "success": True,
                "agents": [
                    {
                        "id": agent.id,
                        "name": agent.name,
                        "description": agent.description,
                        "category": agent.category,
                        "agent_type": agent.agent_type,
                        "variables": agent.variables,
                        "tags": agent.tags,
                        "usage_count": agent.usage_count,
                        "like_count": agent.like_count,
                        "is_official": agent.is_official,
                        "created_at": agent.created_at.isoformat() if agent.created_at else None
                    }
                    for agent in agents
                ]
            }
        finally:
            db.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/agents/{agent_id}")
async def get_agent(agent_id: int):
    """获取智能体详情"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            agent = db.query(Agent).filter(Agent.id == agent_id).first()
            if not agent:
                raise HTTPException(status_code=404, detail="智能体不存在")

            return {
                "success": True,
                "agent": {
                    "id": agent.id,
                    "name": agent.name,
                    "description": agent.description,
                    "category": agent.category,
                    "agent_type": agent.agent_type,
                    "system_prompt": agent.system_prompt,
                    "variables": agent.variables,
                    "ai_model": agent.ai_model,
                    "temperature": agent.temperature,
                    "max_tokens": agent.max_tokens,
                    "batch_count": agent.batch_count,
                    "visibility": agent.visibility,
                    "tags": agent.tags,
                    "usage_count": agent.usage_count,
                    "like_count": agent.like_count,
                    "is_official": agent.is_official,
                    "created_at": agent.created_at.isoformat() if agent.created_at else None
                }
            }
        finally:
            db.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/agents")
async def create_agent(request: dict):
    """创建智能体"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            agent = Agent(
                name=request.get("name"),
                description=request.get("description"),
                category=request.get("category", "custom"),
                agent_type=request.get("agent_type", "user"),
                system_prompt=request.get("system_prompt"),
                variables=request.get("variables", []),
                ai_model=request.get("ai_model", "deepseek"),
                temperature=request.get("temperature", 80),
                max_tokens=request.get("max_tokens", 2048),
                batch_count=request.get("batch_count", 1),
                visibility=request.get("visibility", "private"),
                tags=request.get("tags", []),
                creator_id=request.get("creator_id", ""),
                is_official=request.get("is_official", 0)
            )

            db.add(agent)
            db.commit()
            db.refresh(agent)

            return {
                "success": True,
                "agent_id": agent.id,
                "message": "智能体创建成功"
            }
        finally:
            db.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/agents/{agent_id}")
async def update_agent(agent_id: int, request: dict):
    """更新智能体"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            agent = db.query(Agent).filter(Agent.id == agent_id).first()
            if not agent:
                raise HTTPException(status_code=404, detail="智能体不存在")

            # 更新字段
            if "name" in request:
                agent.name = request["name"]
            if "description" in request:
                agent.description = request["description"]
            if "system_prompt" in request:
                agent.system_prompt = request["system_prompt"]
            if "variables" in request:
                agent.variables = request["variables"]
            if "temperature" in request:
                agent.temperature = request["temperature"]
            if "max_tokens" in request:
                agent.max_tokens = request["max_tokens"]
            if "batch_count" in request:
                agent.batch_count = request["batch_count"]
            if "visibility" in request:
                agent.visibility = request["visibility"]
            if "tags" in request:
                agent.tags = request["tags"]

            agent.updated_at = datetime.now()
            db.commit()

            return {
                "success": True,
                "message": "智能体更新成功"
            }
        finally:
            db.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/agents/{agent_id}")
async def delete_agent(agent_id: int):
    """删除智能体"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            agent = db.query(Agent).filter(Agent.id == agent_id).first()
            if not agent:
                raise HTTPException(status_code=404, detail="智能体不存在")

            db.delete(agent)
            db.commit()

            return {
                "success": True,
                "message": "智能体删除成功"
            }
        finally:
            db.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def replace_variables(prompt: str, variables: Dict) -> str:
    """替换提示词中的变量占位符"""
    import re

    def replacer(match):
        var_name = match.group(1)
        return str(variables.get(var_name, f"{{{{{var_name}}}}}"))

    # 替换 {{变量名}} 格式
    result = re.sub(r'\{\{([^}]+)\}\}', replacer, prompt)
    return result


@app.post("/api/agents/{agent_id}/execute")
async def execute_agent(agent_id: int, request: dict):
    """执行智能体 - 支持批量生成"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            # 获取智能体
            agent = db.query(Agent).filter(Agent.id == agent_id).first()
            if not agent:
                raise HTTPException(status_code=404, detail="智能体不存在")

            # 获取用户输入的变量
            input_variables = request.get("variables", {})
            batch_count = request.get("batch_count", agent.batch_count)

            # 替换提示词中的变量
            final_prompt = replace_variables(agent.system_prompt, input_variables)

            # 创建执行记录
            execution = AgentExecution(
                agent_id=agent.id,
                input_variables=input_variables,
                model_used=agent.ai_model or "deepseek",
                temperature=agent.temperature or 80,
                batch_count=batch_count,
                status="running",
                started_at=datetime.now()
            )
            db.add(execution)
            db.commit()
            db.refresh(execution)

            # 批量生成
            ai_client = DeepSeekClient(api_key=settings.deepseek_api_key, model=settings.deepseek_model)
            versions = []

            for i in range(batch_count):
                try:
                    # 调用AI生成
                    messages = [{"role": "user", "content": final_prompt}]
                    temperature = (agent.temperature or 80) / 100
                    content = ai_client._call_api(messages, temperature=temperature)

                    # 保存生成的版本
                    version = AgentVersion(
                        execution_id=execution.id,
                        agent_id=agent.id,
                        version_number=i + 1,
                        content=content
                    )
                    db.add(version)
                    db.commit()
                    db.refresh(version)

                    versions.append({
                        "version_number": version.version_number,
                        "content": content,
                        "version_id": version.id
                    })
                except Exception as e:
                    print(f"生成版本 {i+1} 失败: {e}")

            # 更新执行状态
            execution.status = "success"
            execution.completed_at = datetime.now()
            db.commit()

            # 更新使用次数
            agent.usage_count = (agent.usage_count or 0) + 1
            db.commit()

            return {
                "success": True,
                "execution_id": execution.id,
                "versions": versions,
                "message": f"成功生成 {len(versions)} 个版本"
            }
        finally:
            db.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/agents/{agent_id}/executions")
async def get_agent_executions(agent_id: int, limit: int = 10):
    """获取智能体的执行历史"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            executions = db.query(AgentExecution)\
                .filter(AgentExecution.agent_id == agent_id)\
                .order_by(AgentExecution.created_at.desc())\
                .limit(limit)\
                .all()

            return {
                "success": True,
                "executions": [
                    {
                        "id": exec.id,
                        "input_variables": exec.input_variables,
                        "status": exec.status,
                        "started_at": exec.started_at.isoformat() if exec.started_at else None,
                        "completed_at": exec.completed_at.isoformat() if exec.completed_at else None
                    }
                    for exec in executions
                ]
            }
        finally:
            db.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/executions/{execution_id}/versions")
async def get_execution_versions(execution_id: int):
    """获取执行记录的所有版本"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            versions = db.query(AgentVersion)\
                .filter(AgentVersion.execution_id == execution_id)\
                .order_by(AgentVersion.version_number.asc())\
                .all()

            return {
                "success": True,
                "versions": [
                    {
                        "id": v.id,
                        "version_number": v.version_number,
                        "content": v.content,
                        "is_selected": v.is_selected,
                        "rating": v.rating
                    }
                    for v in versions
                ]
            }
        finally:
            db.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/versions/{version_id}/select")
async def select_version(version_id: int):
    """选择某个版本（标记为已选中）"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            version = db.query(AgentVersion).filter(AgentVersion.id == version_id).first()
            if not version:
                raise HTTPException(status_code=404, detail="版本不存在")

            # 取消同一次执行的其他版本的选中状态
            db.query(AgentVersion)\
                .filter(AgentVersion.execution_id == version.execution_id)\
                .filter(AgentVersion.id != version_id)\
                .update({"is_selected": 0})

            # 标记当前版本为选中
            version.is_selected = 1
            db.commit()

            return {
                "success": True,
                "message": "版本已选中"
            }
        finally:
            db.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 参考素材管理（写同款功能） ==========

import shutil
from fastapi import UploadFile, File
from typing import Optional
import uuid

# 创建素材上传目录
UPLOAD_DIR = "uploads/materials"
os.makedirs(UPLOAD_DIR, exist_ok=True)


@app.get("/api/materials")
async def list_materials(
    content_type: Optional[str] = None,
    genre: Optional[str] = None,
    tags: Optional[str] = None,
    limit: int = 50
):
    """获取参考素材列表"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            query = db.query(ReferenceMaterial)

            # 筛选条件
            if content_type:
                query = query.filter(ReferenceMaterial.content_type == content_type)
            if genre:
                query = query.filter(ReferenceMaterial.genre == genre)
            if tags:
                # JSON字段搜索（简化处理）
                query = query.filter(ReferenceMaterial.tags.contains(tags.split(',')))

            materials = query.order_by(ReferenceMaterial.created_at.desc()).limit(limit).all()

            return {
                "success": True,
                "materials": [
                    {
                        "id": m.id,
                        "title": m.title,
                        "author": m.author,
                        "source": m.source,
                        "file_type": m.file_type,
                        "content_type": m.content_type,
                        "genre": m.genre,
                        "tags": m.tags,
                        "core_conflict": m.core_conflict,
                        "emotion_style": m.emotion_style,
                        "status": m.status,
                        "usage_count": m.usage_count,
                        "like_count": m.like_count,
                        "is_favorite": m.is_favorite,
                        "created_at": m.created_at.isoformat() if m.created_at else None
                    }
                    for m in materials
                ]
            }
        finally:
            db.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/materials/{material_id}")
async def get_material(material_id: int):
    """获取素材详情"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            material = db.query(ReferenceMaterial).filter(ReferenceMaterial.id == material_id).first()
            if not material:
                raise HTTPException(status_code=404, detail="素材不存在")

            return {
                "success": True,
                "material": {
                    "id": material.id,
                    "title": material.title,
                    "author": material.author,
                    "source": material.source,
                    "source_url": material.source_url,
                    "file_type": material.file_type,
                    "file_size": material.file_size,
                    "content_type": material.content_type,
                    "raw_content": material.raw_content,
                    "analysis": material.analysis,
                    "genre": material.genre,
                    "tags": material.tags,
                    "core_conflict": material.core_conflict,
                    "emotion_style": material.emotion_style,
                    "writing_style": material.writing_style,
                    "characters_extracted": material.characters_extracted,
                    "plot_structure": material.plot_structure,
                    "similarity_tags": material.similarity_tags,
                    "usage_count": material.usage_count,
                    "like_count": material.like_count,
                    "is_favorite": material.is_favorite,
                    "status": material.status,
                    "notes": material.notes,
                    "created_at": material.created_at.isoformat() if material.created_at else None
                }
            }
        finally:
            db.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/materials/upload")
async def upload_material(
    title: str = Form(...),
    author: str = Form(None),
    source: str = Form(None),
    content_type: str = Form(...),
    file: UploadFile = File(...),
    notes: str = Form(None)
):
    """上传参考素材文件"""
    try:
        # 生成唯一文件名
        file_ext = file.filename.split('.')[-1] if '.' in file.filename else 'txt'
        unique_filename = f"{uuid.uuid4()}.{file_ext}"
        file_path = os.path.join(UPLOAD_DIR, unique_filename)

        # 保存文件
        with open(file_path, 'wb') as buffer:
            shutil.copyfileobj(file.file, buffer)
        file_size = os.path.getsize(file_path)

        # 提取文本内容
        raw_content = extract_text_from_file(file_path, file_ext)

        # 创建数据库记录
        db_gen = get_db()
        db = next(db_gen)
        try:
            material = ReferenceMaterial(
                title=title,
                author=author,
                source=source,
                file_type=file_ext,
                file_path=file_path,
                file_size=file_size,
                content_type=content_type,
                raw_content=raw_content[:10000],  # 限制长度
                status="pending",
                uploaded_by="user",
                notes=notes
            )

            db.add(material)
            db.commit()
            db.refresh(material)

            # 异步启动AI分析
            # 注意：这里简化处理，实际应该使用后台任务
            analyze_material_async(material.id)

            return {
                "success": True,
                "material_id": material.id,
                "message": "素材上传成功，正在分析中..."
            }
        finally:
            db.close()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/materials/{material_id}/analyze")
async def analyze_material(material_id: int):
    """手动触发AI分析"""
    try:
        analyze_material_async(material_id)
        return {
            "success": True,
            "message": "分析任务已启动"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/materials/{material_id}/similar")
async def find_similar_materials(material_id: int, limit: int = 5):
    """查找相似的素材（基于标签和题材）"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            target = db.query(ReferenceMaterial).filter(ReferenceMaterial.id == material_id).first()
            if not target:
                raise HTTPException(status_code=404, detail="素材不存在")

            # 简单的相似度推荐（基于题材和标签）
            query = db.query(ReferenceMaterial).filter(ReferenceMaterial.id != material_id)

            if target.genre:
                query = query.filter(ReferenceMaterial.genre == target.genre)

            materials = query.limit(limit).all()

            return {
                "success": True,
                "similar_materials": [
                    {
                        "id": m.id,
                        "title": m.title,
                        "author": m.author,
                        "genre": m.genre,
                        "tags": m.tags,
                        "core_conflict": m.core_conflict
                    }
                    for m in materials
                ]
            }
        finally:
            db.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/materials/{material_id}/write-similar")
async def write_similar(material_id: int, request: dict):
    """基于素材生成同款作品"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            material = db.query(ReferenceMaterial).filter(ReferenceMaterial.id == material_id).first()
            if not material:
                raise HTTPException(status_code=404, detail="素材不存在")

            # 使用AI生成同款
            prompt = f"""你是专业的网文作家。请根据以下参考素材，创作一部同款风格的小说。

**参考作品：**{material.title}
**作者：**{material.author or '未知'}
**题材：**{material.genre or '未知'}
**核心冲突：**{material.core_conflict or '未知'}
**情绪风格：**{material.emotion_style or '未知'}
**写作风格：**{material.writing_style or '未知'}

**参考导语（如果有）：**
{material.raw_content[:500] if material.raw_content else '无'}

**创作要求：**
1. 保持相同的情绪张力和狗血程度
2. 使用不同的情节和人物
3. 保留核心冲突的类型（如金钱、背叛、身份等）
4. 生成300字左右的导语

请直接输出新的导语，不要有任何额外说明："""

            ai_client = DeepSeekClient(api_key=settings.deepseek_api_key, model=settings.deepseek_model)
            messages = [{"role": "user", "content": prompt}]
            content = ai_client._call_api(messages, temperature=0.9)

            # 更新使用次数
            material.usage_count = (material.usage_count or 0) + 1
            db.commit()

            return {
                "success": True,
                "content": content,
                "reference_title": material.title
            }
        finally:
            db.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# 辅助函数

import requests
import time

def extract_pdf_with_insightdoc(file_path: str) -> str:
    """使用 InsightDoc API 解析 PDF 文件（支持 OCR 和复杂格式）"""
    try:
        # 1. 提交任务
        url = f"{settings.insightdoc_base_url}/api/tasks"
        headers = {
            "Authorization": f"Bearer {settings.insightdoc_api_key}"
        }

        with open(file_path, 'rb') as f:
            files = {
                'file': f,
                'task_type': (None, 'docparse'),  # 使用通用文档解析
                'file_name': (None, os.path.basename(file_path))
            }

            response = requests.post(url, headers=headers, files=files, timeout=30)

            if response.status_code != 200:
                print(f"InsightDoc API 错误: {response.status_code} - {response.text}")
                return None

            task_data = response.json()
            task_id = task_data.get('id')

            if not task_id:
                print(f"InsightDoc API 返回无效响应: {task_data}")
                return None

            print(f"✅ InsightDoc 任务已提交: {task_id}")

        # 2. 轮询任务状态
        max_wait = 300  # 最多等待5分钟
        start_time = time.time()
        poll_interval = 2  # 每2秒查询一次

        while time.time() - start_time < max_wait:
            time.sleep(poll_interval)

            status_url = f"{settings.insightdoc_base_url}/api/tasks/detail/{task_id}?result_type=md"
            status_response = requests.get(status_url, headers=headers, timeout=10)

            if status_response.status_code != 200:
                print(f"查询状态失败: {status_response.status_code}")
                time.sleep(poll_interval)
                continue

            status_data = status_response.json()
            status = status_data.get('status')

            if status == 'done':
                print(f"✅ InsightDoc 解析完成")
                # 返回 Markdown 格式的文本
                result = status_data.get('result', '')
                if isinstance(result, dict):
                    # 提取文本内容
                    return result.get('markdown', str(result))
                elif isinstance(result, list):
                    # 处理列表格式
                    return '\n'.join([str(item) for item in result])
                else:
                    return str(result)

            elif status == 'failed':
                error_msg = status_data.get('message', '未知错误')
                print(f"❌ InsightDoc 解析失败: {error_msg}")
                return None

            elif status in ['pending', 'processing']:
                elapsed = int(time.time() - start_time)
                print(f"⏳ InsightDoc 正在解析... ({elapsed}秒)")
            else:
                print(f"⚠️ 未知状态: {status}")

        print(f"⏰ InsightDoc 解析超时")
        return None

    except Exception as e:
        print(f"❌ InsightDoc API 调用失败: {str(e)}")
        return None


def extract_text_from_file(file_path: str, file_ext: str) -> str:
    """从文件中提取文本"""
    try:
        if file_ext.lower() == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_ext.lower() == 'pdf':
            # 优先使用 InsightDoc API（支持 OCR 和复杂格式）
            print(f"📄 使用 InsightDoc API 解析 PDF...")
            insightdoc_result = extract_pdf_with_insightdoc(file_path)

            if insightdoc_result:
                return insightdoc_result
            else:
                print("⚠️ InsightDoc 解析失败，降级使用 PyPDF2...")
                # 降级方案：使用 PyPDF2
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text()
                        return text
                except ImportError:
                    return "PDF解析功能未安装，请运行: pip install PyPDF2"
        elif file_ext.lower() in ['docx', 'doc']:
            # 需要安装 python-docx: pip install python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except ImportError:
                return "Word文档解析功能未安装，请运行: pip install python-docx"
        elif file_ext.lower() in ['xlsx', 'xls']:
            # Excel 文件处理
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                text = ""
                for sheet in wb.worksheets:
                    text += f"\n=== 工作表: {sheet.title} ===\n"
                    for row in sheet.iter_rows(values_only=True):
                        # 过滤空行
                        row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip():
                            text += row_text + "\n"
                return text
            except Exception as e:
                return f"Excel 文件解析失败: {str(e)}"
        else:
            return "不支持的文件格式"
    except Exception as e:
        return f"文件读取失败: {str(e)}"


def analyze_material_async(material_id: int):
    """异步分析素材（简化版本，实际应该用Celery等任务队列）"""
    import threading

    def analyze():
        try:
            db_gen = get_db()
            db = next(db_gen)
            try:
                material = db.query(ReferenceMaterial).filter(ReferenceMaterial.id == material_id).first()
                if not material:
                    return

                # 更新状态
                material.status = "analyzing"
                material.analysis_progress = 10
                db.commit()

                # 使用AI分析素材
                ai_client = DeepSeekClient(api_key=settings.deepseek_api_key, model=settings.deepseek_model)
                prompt = f"""请分析以下小说素材，提取关键信息。以JSON格式返回：

**作品名称：**{material.title}
**内容类型：**{material.content_type}
**原文内容：**
{material.raw_content[:2000]}

请按以下JSON格式返回分析结果：
```json
{{
  "genre": "题材类型（如：世情文、追妻火葬场、豪门恩怨等）",
  "tags": ["标签1", "标签2", "标签3"],
  "core_conflict": "核心冲突描述",
  "emotion_style": "情绪风格（如：压抑、爆发、温暖等）",
  "writing_style": "写作风格描述",
  "characters": [
    {{"name": "角色名", "role": "主角/反派/配角", "traits": ["特征1", "特征2"]}}
  ],
  "plot_points": ["情节点1", "情节点2", "情节点3"]
}}
```

只返回JSON，不要其他文字："""

                messages = [{"role": "user", "content": prompt}]
                response = ai_client._call_api(messages, temperature=0.3)

                # 解析AI返回的JSON
                import json
                import re

                # 尝试提取JSON
                json_match = re.search(r'```json\s*(.*?)\s*```', response, re.DOTALL)
                if json_match:
                    json_str = json_match.group(1)
                else:
                    # 尝试直接找到 {...}
                    start_idx = response.find('{')
                    end_idx = response.rfind('}')
                    if start_idx != -1 and end_idx != -1:
                        json_str = response[start_idx:end_idx + 1]
                    else:
                        raise ValueError("无法提取JSON")

                analysis_result = json.loads(json_str)

                # 更新分析结果
                material.analysis = analysis_result
                material.genre = analysis_result.get("genre", "")
                material.tags = analysis_result.get("tags", [])
                material.core_conflict = analysis_result.get("core_conflict", "")
                material.emotion_style = analysis_result.get("emotion_style", "")
                material.writing_style = analysis_result.get("writing_style", "")
                material.characters_extracted = analysis_result.get("characters", [])
                material.similarity_tags = {tag: 1.0 for tag in analysis_result.get("tags", [])}
                material.status = "completed"
                material.analysis_progress = 100
                material.updated_at = datetime.now()

                db.commit()

                print(f"✅ 素材 {material_id} 分析完成！")
            except Exception as e:
                print(f"❌ 素材分析失败: {e}")
                material.status = "failed"
                material.analysis_progress = 0
                db.commit()
            finally:
                db.close()
        except Exception as e:
            print(f"❌ 分析任务异常: {e}")

    # 在后台线程中执行分析
    thread = threading.Thread(target=analyze)
    thread.daemon = True
    thread.start()


# ========== 初始化预设智能体 ==========

def init_default_agents():
    """初始化默认的智能体模板"""
    try:
        db_gen = get_db()
        db = next(db_gen)
        try:
            # 检查是否已经初始化
            existing = db.query(Agent).filter(Agent.agent_type == "system").count()
            if existing > 0:
                return  # 已经初始化过了

            # 智能体1: 短篇导语生成器
            intro_agent = Agent(
                name="短篇导语生成器",
                description="根据对标作品和导语，生成新的狗血导语",
                category="intro",
                agent_type="system",
                system_prompt="""你是一位专业的网文导语创作专家。请根据以下对标作品的信息，创作3个全新的、狗血的、情绪张力强的导语。

**对标作品：**
{{reference_work}}

**对标导语：**
{{reference_intro}}

**要求：**
1. 保持狗血程度和情绪张力
2. 使用相似的结构但不同的情节
3. 保留核心冲突（金钱、背叛、身份等）
4. 每个导语控制在300字以内
5. 必须有强烈的情绪钩子

请直接输出第1个导语，不要有任何额外说明。""",
                variables=[
                    {"name": "reference_work", "type": "text", "label": "对标作品名称", "default": "", "required": True},
                    {"name": "reference_intro", "type": "textarea", "label": "对标导语内容", "default": "", "required": True}
                ],
                ai_model="deepseek",
                temperature=95,
                max_tokens=1024,
                batch_count=3,
                visibility="public",
                tags=["导语", "短篇", "世情文", "狗血"],
                is_official=1,
                usage_count=0
            )

            # 智能体2: 短篇小说大纲生成器
            outline_agent = Agent(
                name="短篇小说大纲生成器",
                description="根据对标作品和导语，生成完整的短篇大纲",
                category="outline",
                agent_type="system",
                system_prompt="""你是一位专业的网文大纲创作专家。请根据以下信息，生成一个完整的短篇小说大纲。

**对标作品：**
{{reference_work}}

**新导语：**
{{new_intro}}

**要求：**
1. 生成1-2万字短篇大纲
2. 包含8-12个章节
3. 每章有明确的冲突点
4. 情节要狗血、有反转
5. 情绪弧线清晰（压抑-爆发-高潮）

请按以下JSON格式输出：
```json
{
  "title": "小说标题",
  "total_chapters": 10,
  "target_words": 15000,
  "chapters": [
    {
      "chapter_number": 1,
      "title": "章节标题",
      "summary": "本章摘要",
      "plot_points": ["情节点1", "情节点2"],
      "target_words": 1500,
      "emotion_arc": "情绪如何变化"
    }
  ]
}
```""",
                variables=[
                    {"name": "reference_work", "type": "text", "label": "对标作品名称", "default": "", "required": True},
                    {"name": "new_intro", "type": "textarea", "label": "新导语内容", "default": "", "required": True}
                ],
                ai_model="deepseek",
                temperature=90,
                max_tokens=2048,
                batch_count=1,
                visibility="public",
                tags=["大纲", "短篇", "结构"],
                is_official=1,
                usage_count=0
            )

            db.add(intro_agent)
            db.add(outline_agent)
            db.commit()

            print("✅ 默认智能体初始化成功！")
        finally:
            db.close()
    except Exception as e:
        print(f"❌ 初始化默认智能体失败: {e}")


# 启动时初始化默认智能体
init_default_agents()


# ========== 创作灵感助手相关 ==========

class InspirationRequest(BaseModel):
    """灵感生成请求"""
    summary: Optional[str] = None  # 故事简介
    readers: Optional[str] = None  # 目标读者
    genre: Optional[str] = None  # 小说类型
    chapters: Optional[str] = None  # 章节数量
    words: Optional[str] = None  # 目标字数
    elements: Optional[str] = None  # 写作元素


class OutlineRequest(BaseModel):
    """大纲生成请求"""
    settings: Dict  # AI生成的设定


class ChaptersRequest(BaseModel):
    """章节生成请求"""
    settings: Dict  # 设定
    outline: Dict  # 大纲


class NovelRequest(BaseModel):
    """成文请求"""
    settings: Dict  # 设定
    outline: Dict  # 大纲
    chapters: Dict  # 章节


@app.post("/api/inspiration/generate-settings")
async def generate_inspiration_settings(request: InspirationRequest):
    """第一步: 生成小说设定(增强版 - 包含核心矛盾、黄金三章、元素碰撞)"""
    try:
        # 构建提示词 - 一次调用完成所有功能
        prompt = """你是专业的网文创作顾问。请根据用户的灵感,生成完整的小说设定。

**用户输入:**
"""

        if request.summary:
            prompt += f"\n故事简介: {request.summary}"

        if request.readers:
            readers_map = {
                'young': '青少年 (12-18岁)',
                'young_adult': '年轻成人 (18-30岁)',
                'adult': '成人 (30-50岁)',
                'all': '全年龄'
            }
            prompt += f"\n目标读者: {readers_map.get(request.readers, request.readers)}"

        if request.genre:
            genre_map = {
                'romance': '言情小说',
                'fantasy': '玄幻小说',
                'urban': '都市小说',
                'historical': '历史小说',
                'scifi': '科幻小说',
                'mystery': '悬疑小说',
                'wuxia': '武侠小说'
            }
            prompt += f"\n小说类型: {genre_map.get(request.genre, request.genre)}"

        if request.elements:
            prompt += f"\n写作元素: {request.elements}"

        prompt += """

**要求:**
请生成一个简洁但完整的小说设定:

**第一部分: 核心矛盾**
1. 主角的核心欲望
2. 核心阻碍
3. 失败后果
4. 独特卖点

**第二部分: 基本信息**
1. 小说标题 (吸引人,有网感)
2. 副标题
3. 故事简介 (150字左右)

**第三部分: 世界观**
1. 时空背景
2. 世界规则
3. 社会结构

**第四部分: 角色设定** (2-3个主要角色)
对于每个角色,请提供:
- 姓名
- 角色类型 (protagonist/antagonist/supporting)
- 核心身份 (50字)
- 核心性格 (80字,包含缺陷)
- 核心动机 (50字)

**第五部分: 黄金三章**
1. 第一章钩子
2. 第二章冲突
3. 第三章转折

**重要:**
- 设定要简洁有力,避免冗长
- 每个字段控制在50-80字
- 角色描述简明扼要但立体

请严格按照以下JSON格式返回,不要有任何额外文字:
```json
{
  "title": "小说标题",
  "subtitle": "副标题(可选)",
  "summary": "故事简介",
  "core_conflict": {
    "protagonist_desire": "主角想要什么",
    "core_obstacle": "什么在阻碍他",
    "tragic_consequence": "失败的后果",
    "unique_selling_point": "独特卖点"
  },
  "setting": {
    "time_space": "时空背景",
    "world_rules": "世界规则",
    "social_structure": "社会结构"
  },
  "characters": [
    {
      "name": "角色名",
      "role_type": "protagonist",
      "core_identity": "核心身份",
      "core_personality": "核心性格(包含缺陷)",
      "core_motivation": "核心动机"
    }
  ],
  "golden_three_chapters": {
    "chapter1_hook": "第一章钩子",
    "chapter2_conflict": "第二章冲突",
    "chapter3_twist": "第三章转折"
  }
}
```
"""

        # 调用AI生成
        messages = [{"role": "user", "content": prompt}]
        # 降低max_tokens避免API连接问题
        response_content = ai_client._call_api(messages, temperature=0.85, max_tokens=4000, timeout=180.0)

        print("=" * 50)
        print("AI原始响应长度:", len(response_content))
        print("AI原始响应（前2000字符）:")
        print(response_content[:2000])
        print("=" * 50)

        # 使用增强的JSON解析函数
        settings_data = parse_json_response(response_content)

        if not settings_data:
            raise HTTPException(
                status_code=500,
                detail="AI生成的内容格式错误，无法解析为有效的JSON。请重试。"
            )

        return {
            "success": True,
            "data": settings_data
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/inspiration/bubbles")
@app.post("/api/inspiration/bubbles")
async def get_inspiration_bubbles(genre: str = None):
    """获取灵感气泡 - 生成小说的'灵魂' (主角身份+角色设定)

    Args:
        genre: 可选题材参数，如 "甜宠"、"悬疑"、"大女主"、"脑洞"、"末世"、"复仇"、"仙侠"、"都市反转" 等
    """
    try:
        # 题材专属模板 - 不同题材有不同风格的"灵魂"设定
        genre_templates = {
            "甜宠": {
                "theme": "甜宠撒糖",
                "patterns": [
                    "身份反转+A掉马+真香",
                    "误会婚姻+暗中守护+掉马打脸",
                    "替嫁/联姻+隐藏身份+宠妻狂魔",
                    "欢喜冤家+强制绑定+先婚后爱"
                ],
                "examples": [
                    "联姻三年丈夫从未归家，原来他每天在我身边做贴身暗卫。",
                    "说好的替嫁残疾大佬，新婚夜他突然站起来把我抵在墙角。",
                    "高冷上司每晚偷偷溜进我家，说要跟我演一场假戏真做的恋爱。",
                    "被退婚后转身嫁给他小叔，前任在婚礼上红了眼。"
                ]
            },
            "悬疑": {
                "theme": "悬疑惊悚",
                "patterns": [
                    "身份错位+细思极恐的真相",
                    "时间循环+死亡预知",
                    "记忆缺失+身份谜团",
                    "看似正常+内里恐怖"
                ],
                "examples": [
                    "每天醒来都是同一天，而我的丈夫每天都用不同的方式杀我。",
                    "搬进新家后，我发现墙纸下藏着一行字：快逃，他在看着你。",
                    "我失忆后醒来，所有人都叫我'宝贝'，但他们的眼神都在看猎物。",
                    "儿子告诉我床底下有人，我低头一看，那是一张和我一模一样的脸。"
                ]
            },
            "大女主": {
                "theme": "女强爽文",
                "patterns": [
                    "重生复仇+手撕渣男绿茶",
                    "马甲掉落+全员震惊",
                    "穿越逆袭+打脸极品亲戚",
                    "满级大佬+新手村开挂"
                ],
                "examples": [
                    "被渣男退婚后，我转身嫁给他死对头，次日两家股票涨跌分明。",
                    "满级女战神穿成娇滴滴嫡女，拔剑那一刻满朝文武跪了一地。",
                    "首富千金觉醒后收回所有资助，渣哥绿茶姐哭着求我给口饭吃。",
                    "本是修仙界第一人，穿成内卷高考生后我用聚灵阵刷爆了理综卷。"
                ]
            },
            "脑洞": {
                "theme": "脑洞大开",
                "patterns": [
                    "跨物种/跨时空绑定",
                    "系统故障+奇葩任务",
                    "物种交换+人类行为",
                    "打破第四面墙"
                ],
                "examples": [
                    "我的猫每晚变成人，教我怎么在职场做人，他曾是上市公司CEO。",
                    "绑定'反派洗白系统'后，我被迫攻略所有被我伤害过的男主。",
                    "穿越成手机，每天都要帮机主打字回复男神消息，但他怎么越来越暧昧？",
                    "能听到物品说话后，我家菜刀哭着说：今晚你老公又要用它表演劈叉了。"
                ]
            },
            "末世": {
                "theme": "末世求生",
                "patterns": [
                    "重生囤货+空间开挂",
                    "异能觉醒+反派求饶",
                    "基地建设+收服小弟",
                    "背叛复仇+独自称王"
                ],
                "examples": [
                    "重生回末世前三天，我卖掉公司买下全国最大仓储基地，建起钢铁堡垒。",
                    "末世觉醒SSS级治愈异能，渣男前夫跪求我救命，我转身喂了丧尸。",
                    "别人都在逃命，我却带着满级空间在末世开超市，丧尸也要排队扫码。",
                    "被队友推进丧尸堆后觉醒，我成了丧尸女王，回头问他们还跑吗。"
                ]
            },
            "复仇": {
                "theme": "复仇爽文",
                "patterns": [
                    "全家灭门+十年归来",
                    "被害惨死+重生改命",
                    "身份互换+以牙还牙",
                    "温柔刀+杀人诛心"
                ],
                "examples": [
                    "全家被灭门的那个雨夜，我躲在衣柜里，记住了每一个人的脸。",
                    "被渣男闺蜜害死后，我重生回订婚宴，当着所有人的手撕了他们。",
                    "十年后我以新身份归来，曾经的霸凌者跪在地上求我认出他。",
                    "养父杀父夺母之仇，我用了二十年，叫他一声父，送他下地狱。"
                ]
            },
            "仙侠": {
                "theme": "仙侠玄幻",
                "patterns": [
                    "废材逆袭+血脉觉醒",
                    "师徒禁忌+身份反转",
                    "魔尊/仙尊掉马+追妻火葬",
                    "穿书+改变宿命"
                ],
                "examples": [
                    "废材二小姐觉醒凤凰血脉，那三宗四门跪着求我收他们为徒。",
                    "仙尊师父下凡渡劫，我陪他当了十年凡人妻，他飞升后却忘了我是谁。",
                    "穿成恶毒女配后，我把男主全杀了，魔尊却说要嫁给我做压寨夫人。",
                    "本是修仙界第一剑仙，下山第一件事是去豪门当保姆，只因雇主长得像他。"
                ]
            },
            "都市反转": {
                "theme": "都市爽文",
                "patterns": [
                    "隐藏身份+掉马打脸",
                    "装穷装废+惊艳反转",
                    "被看不起+实力碾压",
                    "替身反转+真爱反转"
                ],
                "examples": [
                    "相亲男嫌弃我是收银员，不知道这家超市是我生日礼物。",
                    "我装了三年穷屌丝，女朋友提分手那天，我开着法拉利来搬家。",
                    "全校都嘲笑我是低保户，直到有一天几十辆豪车来接我回家继承家业。",
                    "假千金炫耀她未婚夫是总裁，而总裁正在我家给我爸端茶倒水。"
                ]
            }
        }

        import random

        # 如果指定了题材，使用该题材模板；否则随机选择
        if genre and genre in genre_templates:
            template = genre_templates[genre]
            selected_genres = [genre]
        else:
            genres_list = list(genre_templates.keys())
            selected_genres = random.sample(genres_list, 3)
            genre = random.choice(selected_genres)
            template = genre_templates[genre]

        # 构建增强的 prompt
        prompt = f"""你是小说灵感大师，专门生成小说的「灵魂设定」。

**核心原则：** 一篇小说的灵魂=主角的身份/角色设定。这个设定必须足够奇特、有反差、有情绪、有爽点，让读者一眼就想知道"这会是一个怎样的故事"。

**当前题材：** {genre}
**核心风格：** {template['theme']}

**该题材的灵魂模式：**
{chr(10).join(f"- {p}" for p in template['patterns'])}

**该题材参考案例（理解这种感觉，但不要照抄）：**
{chr(10).join(f"- {e}" for e in template['examples'])}

**请生成8条全新的、充满「灵魂」的灵感短句，要求：**

1. **聚焦主角设定** - 每条都要围绕主角的身份、角色、处境展开
2. **强反差/强冲突** - 身份的反转、地位的落差、认知的颠覆
3. **有情绪有爽点** - 让人期待"接下来会发生什么"
4. **画面感极强** - 一句话就能让人脑补出完整场景
5. **长度控制在18-35字** - 太短说不清楚，太长失去冲击力
6. **不要和参考案例重复** - 创造全新的设定

请以JSON数组格式返回，不要有任何额外文字：
```json
[
  "灵感1...",
  "灵感2...",
  "..."
]
```
"""

        messages = [{"role": "user", "content": prompt}]
        response_content = ai_client._call_api(messages, temperature=0.95, max_tokens=2000, timeout=60.0)
        
        # 解析JSON
        try:
            # 提取JSON部分
            start_idx = response_content.find('[')
            end_idx = response_content.rfind(']')
            if start_idx != -1 and end_idx != -1:
                json_str = response_content[start_idx:end_idx + 1]
                inspirations = json.loads(json_str)
            else:
                # 备选解析逻辑
                inspirations = parse_json_response(response_content)
                if not isinstance(inspirations, list):
                    inspirations = []
        except Exception as e:
            print(f"解析灵感气泡JSON失败: {e}")
            inspirations = []

        # 保底数据 - 按题材分类
        fallback_inspirations = {
            "甜宠": [
                "联姻三年丈夫从未归家，原来他每天在我身边做贴身暗卫。",
                "说好的替嫁残疾大佬，新婚夜他突然站起来把我抵在墙角。",
                "高冷上司每晚偷偷溜进我家，说要跟我演一场假戏真做的恋爱。",
                "被退婚后转身嫁给他小叔，前任在婚礼上红了眼。",
                "相亲对象嫌弃我太胖，他不知道我是他爱吃的胖胖面包店老板。",
                "死对头成了我的合租室友，每天早上还要抢我的牙膏。",
                "暗恋十年的男神向我求婚，我高兴了一整晚，他却在婚礼上说谢谢我不嫌弃他。",
                "收养的流浪猫每晚变成人，教我怎么攻略高冷男神。"
            ],
            "悬疑": [
                "每天醒来都是同一天，而我的丈夫每天都用不同的方式杀我。",
                "搬进新家后，我发现墙纸下藏着一行字：快逃，他在看着你。",
                "我失忆后醒来，所有人都叫我'宝贝'，但他们的眼神都在看猎物。",
                "儿子告诉我床底下有人，我低头一看，那是一张和我一模一样的脸。",
                "能听到别人的心声后，我发现我老公一直在思考怎么把我做成标本。",
                "我的日记本里写着明天的日期，而明天我真的死了一次。",
                "全家都说我疯了，可我明明看见他们把我的尸体埋在后院。",
                "每晚12点我家都会多出一个人，第二天早上又消失不见。"
            ],
            "大女主": [
                "被渣男退婚后，我转身嫁给他死对头，次日两家股票涨跌分明。",
                "满级女战神穿成娇滴滴嫡女，拔剑那一刻满朝文武跪了一地。",
                "首富千金觉醒后收回所有资助，渣哥绿茶姐哭着求我给口饭吃。",
                "本是修仙界第一人，穿成内卷高考生后我用聚灵阵刷爆了理综卷。",
                "穿成虐文女主后，我把剧情全部改成了极简风：能动手绝不吵架。",
                "假千金炫耀她未婚夫是总裁，而总裁正在我家给我爸端茶倒水。",
                "全校都嘲笑我是低保户，直到有一天几十辆豪车来接我回家继承家业。",
                "被推进丧尸堆后觉醒，我成了丧尸女王，回头问队友还跑吗。"
            ],
            "脑洞": [
                "我的猫每晚变成人，教我怎么在职场做人，他曾是上市公司CEO。",
                "绑定'反派洗白系统'后，我被迫攻略所有被我伤害过的男主。",
                "穿越成手机，每天都要帮机主打字回复男神消息，但他怎么越来越暧昧？",
                "能听到物品说话后，我家菜刀哭着说：今晚你老公又要用它表演劈叉了。",
                "世家富贵男主魂移贪财女主身体里，与她共用躯体替她开挂攀高枝。",
                "每天都能听到未来儿子的心声，他告诉我一定要离那个黑脸反派远一点。",
                "修仙大佬穿成内卷高考生，直接用聚灵阵在考场刷题。",
                "明明是九代单传的剑仙，下山第一件事竟然是去应聘豪门保姆。"
            ],
            "末世": [
                "重生回末世前三天，我卖掉公司买下全国最大仓储基地，建起钢铁堡垒。",
                "末世觉醒SSS级治愈异能，渣男前夫跪求我救命，我转身喂了丧尸。",
                "别人都在逃命，我却带着满级空间在末世开超市，丧尸也要排队扫码。",
                "被队友推进丧尸堆后觉醒，我成了丧尸女王，回头问他们还跑吗。",
                "满级火系异能者重生回末世前三天，她买下了全国最优秀的仓储避风港。",
                "末世降临前，我囤了一仓库方便面，成了全人类的救世主。",
                "别人异能是攻击治愈，我的异能是让丧尸乖乖排队做核酸检测。",
                "重生后我第一时间杀了自己的未婚夫，因为上一世是他把我推向了丧尸。"
            ],
            "复仇": [
                "全家被灭门的那个雨夜，我躲在衣柜里，记住了每一个人的脸。",
                "被渣男闺蜜害死后，我重生回订婚宴，当着所有人的手撕了他们。",
                "十年后我以新身份归来，曾经的霸凌者跪在地上求我认出他。",
                "养父杀父夺母之仇，我用了二十年，叫他一声父，送他下地狱。",
                "被推进丧尸堆后觉醒，我成了丧尸女王，回头让队友排队受死。",
                "全家被灭门时我躲在地窖，十年后我以杀手身份归来血洗仇府。",
                "前世被渣男害死，重生后我把他的白月光送到了他床上，让他身败名裂。",
                "我用了十年时间成为他们的雇主，每天以折磨他们为乐。"
            ],
            "仙侠": [
                "废材二小姐觉醒凤凰血脉，那三宗四门跪着求我收他们为徒。",
                "仙尊师父下凡渡劫，我陪他当了十年凡人妻，他飞升后却忘了我是谁。",
                "穿成恶毒女配后，我把男主全杀了，魔尊却说要嫁给我做压寨夫人。",
                "本是修仙界第一剑仙，下山第一件事是去豪门当保姆，只因雇主长得像他。",
                "明明是九代单传的剑仙，下山第一件事竟然是去应聘豪门保姆。",
                "说好的替嫁残疾大佬，谁知道新婚夜他突然站起来教我炼丹。",
                "修仙大佬穿成内卷高考生，直接用聚灵阵在考场刷题。",
                "本想安静做个纨绔王妃，却发现府里的影卫全是我的前世旧部。"
            ],
            "都市反转": [
                "相亲男嫌弃我是收银员，不知道这家超市是我生日礼物。",
                "我装了三年穷屌丝，女朋友提分手那天，我开着法拉利来搬家。",
                "全校都嘲笑我是低保户，直到有一天几十辆豪车来接我回家继承家业。",
                "假千金炫耀她未婚夫是总裁，而总裁正在我家给我爸端茶倒水。",
                "被公司裁员后，我转身买下公司，前老板跪着求我不要开除他。",
                "我装了三年穷屌丝，女友提分手那天，我开着法拉利来搬家。",
                "相亲对象嫌弃我是收银员，却不知道这家超市是我生日礼物。",
                "全家都看不起我摆地摊，直到一辆辆劳斯莱斯来我的摊位进货。"
            ]
        }

        # 如果生成失败，返回对应题材的保底数据
        if not inspirations:
            inspirations = fallback_inspirations.get(genre, fallback_inspirations["脑洞"])

        return {
            "success": True,
            "data": inspirations,
            "genre": genre
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        # 发生错误时也返回保底数据，确保前端不崩溃
        fallback_data = [
            "世家富贵男主魂移贪财女主身体里，共用躯体开挂攀高枝。",
            "明明是九代单传的剑仙，下山第一件事竟然是去应聘豪门保姆。",
            "穿成虐文女主后，我把剧情全部改成了极简风：能动手绝不吵架。"
        ]
        return {
            "success": True,
            "data": fallback_data,
            "genre": genre or "脑洞"
        }


@app.post("/api/inspiration/generate-outline")
async def generate_inspiration_outline(request: OutlineRequest):
    """第二步: 生成章节大纲(增强版 - 包含动态节拍器、伏笔预埋、情感曲线)"""
    try:
        settings = request.settings

        # 确定章节数
        chapter_count = 10  # 默认值

        # 计算动态节拍器(字数分配)
        total_words = 30000
        word_distribution = {
            "opening": int(total_words * 0.15),    # 开篇 15%
            "setup": int(total_words * 0.10),      # 起步 10%
            "conflict": int(total_words * 0.40),   # 中段冲突 40%
            "climax": int(total_words * 0.20),     # 高潮 20%
            "ending": int(total_words * 0.15)      # 结尾 15%
        }

        # 构建提示词
        prompt = f"""你是专业的网文大纲创作专家。请根据以下小说设定,生成详细的章节大纲。

**小说信息:**
标题: {settings.get('title', '未命名')}
简介: {settings.get('summary', '')}

**核心矛盾:**
{json.dumps(settings.get('core_conflict', {}), ensure_ascii=False, indent=2)}

**世界观:**
时空背景: {settings.get('setting', {}).get('time_space', '')}
社会结构: {settings.get('setting', {}).get('social_structure', '')}

**黄金三章锚点:**
第一章钩子: {settings.get('golden_three_chapters', {}).get('chapter1_hook', '')}
第二章冲突: {settings.get('golden_three_chapters', {}).get('chapter2_conflict', '')}
第三章转折: {settings.get('golden_three_chapters', {}).get('chapter3_twist', '')}

**主要角色:**
"""

        # 添加角色信息
        for char in settings.get('characters', [])[:5]:
            role_label = "主角" if char.get('role_type') == 'protagonist' else "反派" if char.get('role_type') == 'antagonist' else "配角"
            prompt += f"\n- {char.get('name', '')} ({role_label}): {char.get('core_identity', '')}\n  性格缺陷: {char.get('personality_flaw', 'N/A')}\n"

        prompt += f"""

**动态节拍器 (字数分配):**
- 开篇 (第1-2章): {word_distribution['opening']}字 - 建立世界观,抛出钩子
- 起步 (第3-4章): {word_distribution['setup']}字 - 角色入场,初步冲突
- 中段冲突 (第5-7章): {word_distribution['conflict']}字 - 矛盾升级,困境重重
- 高潮 (第8-9章): {word_distribution['climax']}字 - 核心冲突爆发
- 结尾 (第10章): {word_distribution['ending']}字 - 收尾,反转,余韵

**要求:**
请生成 {chapter_count} 章的详细大纲,要求:

1. 每章包含:
   - 章节号 (chapter_number)
   - 章节标题 (title) - 要有吸引力
   - 章节摘要 (summary) - 150-200字
   - 情节要点 (plot_points) - 3-5个关键情节
   - 目标字数 (target_words) - 严格遵循动态节拍器分配
   - 涉及角色 (characters) - 参与本章的主要角色
   - 情绪强度 (emotion_intensity) - 1-10,表示本章情绪张力
   - 情绪类型 (emotion_type) - 压抑/紧张/爆发/温馨/绝望等

2. 伏笔预埋:
   - 在合适的位置标记伏笔 (foreshadowing)
   - 标记伏笔回收章节 (resolve_chapter)

3. 情感曲线:
   - 确保情绪有起伏,不要平淡
   - 开篇要有钩子,中段要有张力,高潮要有爆发力

4. 角色性格缺陷体现:
   - 在大纲中标记角色缺陷如何影响剧情

5. 钩子设计:
   - 每章结尾要有悬念或转折
   - 确保读者想继续看下一章

请严格按照以下JSON格式返回:
```json
{{
  "total_chapters": {chapter_count},
  "target_words": {total_words},
  "word_distribution": {json.dumps(word_distribution)},
  "emotion_curve": [
    {{"chapter": 1, "intensity": 7, "type": "紧张"}},
    {{"chapter": 2, "intensity": 5, "type": "压抑"}}
  ],
  "foreshadowing_map": [
    {{"hint": "伏笔内容", "planted_chapter": 1, "resolve_chapter": 7}}
  ],
  "chapters": [
    {{
      "chapter_number": 1,
      "title": "章节标题",
      "summary": "章节摘要",
      "plot_points": ["情节1", "情节2", "情节3"],
      "target_words": 3000,
      "characters": ["角色名1", "角色名2"],
      "emotion_intensity": 7,
      "emotion_type": "紧张",
      "chapter_hook": "结尾钩子",
      "flaw_manifestation": "本章体现的角色缺陷",
      "foreshadowing": ["伏笔1"]
    }}
  ]
}}
```
"""

        # 调用AI生成
        messages = [{"role": "user", "content": prompt}]
        response_content = ai_client._call_api(messages, temperature=0.8, max_tokens=8000, timeout=600)  # 增加到8000 tokens，超时10分钟

        # 打印原始响应用于调试
        print("=" * 50)
        print("AI原始响应长度:", len(response_content))
        print("AI原始响应（前3000字符）:")
        print(response_content[:3000])
        print("=" * 50)

        # 解析JSON
        outline_data = parse_json_response(response_content)

        if outline_data is None:
            # 返回原始内容作为fallback
            return {
                "success": True,
                "data": {
                    "raw_response": response_content,
                    "parse_error": "JSON解析失败，返回原始内容"
                }
            }

        return {
            "success": True,
            "data": outline_data
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


def fix_truncated_json(json_str: str) -> str:
    """尝试修复被截断的JSON字符串"""
    import re

    # 检查是否被截断（最后一个字符不是闭合括号）
    if not json_str.rstrip().endswith('}') and not json_str.rstrip().endswith(']'):
        # 跟踪括号平衡和在字符串内部的状态
        brace_count = 0
        bracket_count = 0
        in_string = False
        escape_next = False
        last_complete_pos = -1
        last_field_end_pos = -1

        for i, char in enumerate(json_str):
            if escape_next:
                escape_next = False
                continue

            if char == '\\':
                escape_next = True
                continue

            if char == '"' and not escape_next:
                in_string = not in_string
                # 记录字符串结束的位置
                if not in_string:
                    last_field_end_pos = i
                continue

            # 如果在字符串内部，跳过括号计数
            if in_string:
                continue

            if char == '{':
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                # 当括号平衡且不在字符串内时，记录位置
                if brace_count == 0 and bracket_count == 0:
                    last_complete_pos = i
            elif char == '[':
                bracket_count += 1
            elif char == ']':
                bracket_count -= 1
                # 当括号平衡且不在字符串内时，记录位置
                if brace_count == 0 and bracket_count == 0:
                    last_complete_pos = i

        # 如果找到了完整的位置，截取到那里
        if last_complete_pos > 0:
            json_str = json_str[:last_complete_pos + 1]
            return json_str

        # 如果没有完整的位置，回退到上一个字段的结尾
        if last_field_end_pos > 0:
            # 找到上一个字段结尾后的位置
            # 需要找到字段的结束符号（逗号或括号）
            truncated = json_str[:last_field_end_pos + 1]

            # 查找下一个非空白字符
            j = last_field_end_pos + 1
            while j < len(json_str) and json_str[j].isspace():
                j += 1

            # 如果后面是逗号，包含它
            if j < len(json_str) and json_str[j] == ',':
                truncated = json_str[:j + 1]
            else:
                # 没有逗号，添加一个
                truncated += ','

            # 闭合括号
            open_braces = truncated.count('{') - truncated.count('}')
            open_brackets = truncated.count('[') - truncated.count(']')

            truncated += '}' * open_braces
            truncated += ']' * open_brackets

            return truncated

        # 最后的手段：简单地闭合括号
        if json_str.count('"') % 2 != 0:
            json_str += '"'

        open_braces = json_str.count('{') - json_str.count('}')
        open_brackets = json_str.count('[') - json_str.count(']')

        json_str += '}' * open_braces
        json_str += ']' * open_brackets

    return json_str


def parse_json_response(response_content: str):
    """智能解析AI返回的JSON，包含多层容错和修复"""
    import json
    import re

    # 尝试多种方式提取JSON
    json_str = None

    # 方法1: 从代码块中提取
    json_match = re.search(r'```json\s*(.*?)\s*```', response_content, re.DOTALL)
    if json_match:
        json_str = json_match.group(1)
    else:
        # 方法2: 找到最外层的{}
        start_idx = response_content.find('{')
        end_idx = response_content.rfind('}') + 1
        if start_idx != -1 and end_idx > start_idx:
            json_str = response_content[start_idx:end_idx]

    if not json_str:
        return None

    # 尝试修复截断的JSON
    json_str = fix_truncated_json(json_str)

    # 尝试直接解析
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        print(f"首次JSON解析失败: {e}")

    # 尝试清理后解析
    json_str_clean = clean_json_string(json_str)
    try:
        return json.loads(json_str_clean)
    except json.JSONDecodeError as e:
        print(f"清理后JSON解析失败: {e}")

    # 尝试修复JSON
    json_str_fixed = fix_json_string(json_str_clean)
    try:
        return json.loads(json_str_fixed)
    except json.JSONDecodeError as e:
        print(f"修复后JSON解析失败: {e}")
        return None


def clean_json_string(json_str: str) -> str:
    """清理JSON字符串中的常见问题"""
    import re

    # 移除注释
    json_str = re.sub(r'//.*?\n', '\n', json_str)
    json_str = re.sub(r'/\*.*?\*/', '', json_str, flags=re.DOTALL)

    # 移除控制字符（除了换行、制表符等）
    json_str = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', json_str)

    # 移除多余的逗号（如 {, } 或 [, ]）
    json_str = re.sub(r',\s*([}\]])', r'\1', json_str)

    return json_str


def fix_json_string(json_str: str) -> str:
    """尝试修复JSON字符串中的常见错误"""
    import re

    # 移除所有换行和多余空格，便于处理
    json_str = re.sub(r'\n+', ' ', json_str)
    json_str = re.sub(r'\s+', ' ', json_str)

    # 尝试修复常见的格式问题
    # 1. 确保对象属性名有引号
    json_str = re.sub(r'([{,]\s*)([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1"\2":', json_str)

    # 2. 修复单引号问题（JSON只接受双引号）
    json_str = json_str.replace("'", '"')

    # 3. 修复未转义的换行符在字符串中
    # 这个比较复杂，暂时跳过

    return json_str


@app.post("/api/inspiration/generate-chapters")
async def generate_inspiration_chapters(request: ChaptersRequest):
    """第三步: 生成章节内容"""
    try:
        print("=" * 50)
        print("收到生成章节请求")
        print("=" * 50)

        settings = request.settings
        outline = request.outline

        print(f"大纲章节数: {len(outline.get('chapters', []))}")

        # 检查是否有解析错误
        if outline.get('parse_error'):
            raise HTTPException(
                status_code=400,
                detail="上一步的大纲生成失败，请重新生成大纲。原始响应：" + outline.get('raw_response', '')[:500]
            )

        chapters = outline.get('chapters', [])

        print(f"大纲包含 {len(chapters)} 个章节")

        # 验证章节数据完整性
        valid_chapters = []
        for idx, ch in enumerate(chapters):
            if not ch.get('summary'):
                print(f"警告: 第 {idx+1} 章缺少 summary 字段，跳过")
                continue
            valid_chapters.append(ch)

        print(f"有效的章节数: {len(valid_chapters)}")

        if not valid_chapters:
            # 检查是否有raw_response
            if 'raw_response' in outline:
                raise HTTPException(
                    status_code=400,
                    detail="大纲格式错误：AI返回的不是有效的大纲格式。请重新生成大纲。"
                )
            else:
                raise HTTPException(status_code=400, detail="大纲中没有有效章节，请检查上一步的大纲生成结果")

        # 使用验证后的章节列表
        chapters = valid_chapters

        # 构建角色上下文
        character_context = "\n".join([
            f"**{char.get('name', '')}** ({char.get('role_type', '')}): {char.get('core_identity', '')}\n  性格: {char.get('core_personality', '')}\n  动机: {char.get('core_motivation', '')}"
            for char in settings.get('characters', [])
        ])

        # 生成每个章节的内容
        generated_chapters = []

        print(f"开始生成 {len(chapters)} 个章节...")

        # 定义单章生成函数
        def generate_single_chapter_task(idx, chapter_outline):
            chapter_num = chapter_outline.get('chapter_number', idx + 1)
            title = chapter_outline.get('title', f'第{chapter_num}章')
            summary = chapter_outline.get('summary', '')
            plot_points = chapter_outline.get('plot_points', [])
            target_words = chapter_outline.get('target_words', 2000)
            characters_involved = chapter_outline.get('characters', [])

            print(f"正在生成第 {idx+1}/{len(chapters)} 章: {title}")

            try:
                # 构建章节生成提示词
                prompt = f"""你是专业的网文作家。请根据以下大纲创作章节内容。

**小说信息:**
标题: {settings.get('title', '未命名')}
本章: 第{chapter_num}章 - {title}

**本章摘要:**
{summary}

**情节要点:**
{chr(10).join([f'- {p}' for p in plot_points])}

**涉及角色:**
{', '.join(characters_involved) if characters_involved else '未指定'}

**角色设定:**
{character_context}

**创作要求:**
1. 严格按照大纲展开情节
2. 语言流畅,符合网文风格
3. 对话要符合角色性格
4. 有情绪张力和冲突
5. 字数约 {target_words} 字
6. 情节要紧凑,不要拖沓
7. 结尾要有悬念或转折

请直接输出章节内容,不要有任何额外说明或标题。
"""

                # 调用AI生成
                messages = [{"role": "user", "content": prompt}]
                chapter_content = ai_client._call_api(messages, temperature=0.85, max_tokens=4000)

                print(f"第 {idx+1} 章生成完成，内容长度: {len(chapter_content)}")

                return {
                    "chapter_number": chapter_num,
                    "title": title,
                    "summary": summary,
                    "content": chapter_content,
                    "word_count": len(chapter_content),
                    "plot_points": plot_points,
                    "characters": characters_involved,
                    "index": idx  # 用于排序
                }
            except Exception as e:
                print(f"第 {idx+1} 章生成失败: {e}")
                import traceback
                traceback.print_exc()
                # 返回失败信息，不抛出异常以免影响其他线程
                return None

        # 并行生成
        print(f"开始并行生成 {len(chapters)} 个章节...")
        generated_chapters = []
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            # 提交所有任务
            future_to_idx = {
                executor.submit(generate_single_chapter_task, idx, ch): idx 
                for idx, ch in enumerate(chapters)
            }
            
            # 收集结果
            results = []
            for future in concurrent.futures.as_completed(future_to_idx):
                result = future.result()
                if result:
                    results.append(result)
            
            # 按原始顺序排序
            results.sort(key=lambda x: x["index"])
            
            # 移除index字段并添加到最终列表
            for res in results:
                del res["index"]
                generated_chapters.append(res)

        print(f"章节生成完成，成功生成 {len(generated_chapters)}/{len(chapters)} 章")

        if not generated_chapters:
            raise HTTPException(
                status_code=500,
                detail="所有章节生成均失败，请检查AI服务配置或稍后重试"
            )

        return {
            "success": True,
            "data": {
                "chapters": generated_chapters
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/inspiration/generate-novel")
async def generate_inspiration_novel(request: NovelRequest):
    """第四步: 一键成文(创建项目)"""
    try:
        settings = request.settings
        outline = request.outline
        chapters_data = request.chapters

        # 创建数据库记录
        db = next(get_db())

        # 创建项目
        project = NovelProject(
            name=settings.get('title', '未命名'),
            theme=settings.get('summary', ''),
            background=settings.get('setting', {}).get('time_space', ''),
            genre=settings.get('subtitle', ''),
            core_conflict=settings.get('summary', ''),
            target_words=outline.get('target_words', 30000),
            status="completed",
            outline=outline,
            characters=settings.get('characters', []),
            chapters=[],
            word_count=0
        )

        db.add(project)
        db.commit()
        db.refresh(project)

        # 转换章节格式
        chapters = []
        total_words = 0

        for ch in chapters_data.get('chapters', []):
            chapter = {
                "id": ch.get('chapter_number', 1),
                "title": ch.get('title', ''),
                "summary": ch.get('summary', ''),
                "content": ch.get('content', ''),
                "word_count": ch.get('word_count', 0),
                "order": ch.get('chapter_number', 1)
            }
            chapters.append(chapter)
            total_words += ch.get('word_count', 0)

        # 更新项目
        project.chapters = chapters
        project.word_count = total_words
        db.commit()

        # ==========================================
        # 核心逻辑修复：填充规范化数据库表
        # ==========================================
        
        # 1. 创建角色表数据
        for char in settings.get('characters', []):
            try:
                character = Character(
                    project_id=project.id,
                    name=char.get('name', '未命名'),
                    role_type=char.get('role_type', 'supporting'),
                    importance='core' if char.get('role_type') == 'protagonist' else 'important',
                    core_identity=char.get('core_identity'),
                    core_personality=char.get('core_personality'),
                    core_motivation=char.get('core_motivation'),
                    personality_flaw=char.get('personality_flaw'),
                    flaw_consequence=char.get('flaw_consequence'),
                    growth_direction=char.get('growth_direction'),
                    source="ai_generated"
                )
                db.add(character)
            except Exception as e:
                print(f"创建角色失败: {e}")
        
        # 2. 创建大纲和章节草稿
        for idx, ch in enumerate(chapters_data.get('chapters', [])):
            try:
                # 匹配大纲信息 (尝试从outline中找对应章节)
                outline_info = {}
                for out_ch in outline.get('chapters', []):
                    if out_ch.get('chapter_number') == ch.get('chapter_number'):
                        outline_info = out_ch
                        break
                
                # 创建大纲记录
                plot_outline = PlotOutline(
                    project_id=project.id,
                    level="chapter",
                    chapter_number=ch.get('chapter_number', idx + 1),
                    title=ch.get('title', ''),
                    summary=ch.get('summary', '') or outline_info.get('summary', ''),
                    plot_points=outline_info.get('plot_points', []),
                    target_words=outline_info.get('target_words', 3000),
                    focus_elements=[],
                    emotion_arc=f"{outline_info.get('emotion_type', '')} {outline_info.get('emotion_intensity', '')}",
                    characters_involved=outline_info.get('characters', []),
                    source="ai_generated",
                    status="generated",
                    order=ch.get('chapter_number', idx + 1)
                )
                db.add(plot_outline)
                db.flush() # 获取ID
                
                # 创建章节草稿
                chapter_draft = ChapterDraft(
                    project_id=project.id,
                    outline_id=plot_outline.id,
                    chapter_number=ch.get('chapter_number', idx + 1),
                    title=ch.get('title', ''),
                    content=ch.get('content', ''),
                    word_count=ch.get('word_count', 0),
                    status="completed",
                    edit_count=0,
                    ai_revision_count=1,
                    human_ai_ratio="0:100",
                    generation_params={"source": "inspiration_assistant"}
                )
                db.add(chapter_draft)
                
            except Exception as e:
                print(f"创建章节数据失败 (章节 {idx+1}): {e}")
        
        db.commit()

        # 保存到稿件表 (Manuscript)
        try:
            manuscript = Manuscript(
                project_id=project.id,
                title=project.name,
                content=chapters,
                created_at=datetime.now(),
                updated_at=datetime.now()
            )
            db.add(manuscript)
            db.commit()
            print(f"小说 {project.id} 稿件记录已保存")
        except Exception as ms_err:
            print(f"保存稿件记录失败: {ms_err}")

        return {
            "success": True,
            "data": {
                "project_id": project.id,
                "title": project.name,
                "chapters": chapters,
                "total_words": total_words
            },
            "message": "小说生成完成!"
        }


    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# 稿件记录 API (Manuscript History)
# =====================================================================

@app.get("/api/manuscripts")
async def list_manuscripts():
    """获取稿件历史列表"""
    try:
        db = next(get_db())
        manuscripts = db.query(Manuscript).order_by(Manuscript.created_at.desc()).all()
        
        result = []
        for m in manuscripts:
            result.append({
                "id": m.id,
                "title": m.title,
                "status": m.status,
                "grade": m.grade,
                "project_id": m.project_id,
                "created_at": m.created_at.strftime("%Y-%m-%d %H:%M:%S")
            })
            
        return {"success": True, "data": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/manuscripts/{manuscript_id}")
async def get_manuscript(manuscript_id: int):
    """获取稿件详情 (包含步骤信息)"""
    try:
        db = next(get_db())
        manuscript = db.query(Manuscript).filter(Manuscript.id == manuscript_id).first()
        
        if not manuscript:
            raise HTTPException(status_code=404, detail="稿件不存在")
            
        steps = db.query(ManuscriptStep).filter(
            ManuscriptStep.manuscript_id == manuscript_id
        ).order_by(ManuscriptStep.created_at).all()
        
        steps_data = []
        for s in steps:
            steps_data.append({
                "id": s.id,
                "step_name": s.step_name,
                "created_at": s.created_at.strftime("%Y-%m-%d %H:%M:%S")
            })
            
        return {
            "success": True, 
            "data": {
                "id": manuscript.id,
                "title": manuscript.title,
                "content": manuscript.content,
                "review_report": manuscript.review_report,
                "grade": manuscript.grade,
                "status": manuscript.status,
                "project_id": manuscript.project_id,
                "created_at": manuscript.created_at.strftime("%Y-%m-%d %H:%M:%S"),
                "steps": steps_data
            }
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/manuscripts/{manuscript_id}/steps/{step_name}")
async def get_manuscript_step(manuscript_id: int, step_name: str):
    """获取特定步骤的中间产物"""
    try:
        db = next(get_db())
        step = db.query(ManuscriptStep).filter(
            ManuscriptStep.manuscript_id == manuscript_id,
            ManuscriptStep.step_name == step_name
        ).order_by(ManuscriptStep.created_at.desc()).first()
        
        if not step:
            raise HTTPException(status_code=404, detail=f"找不到步骤 {step_name} 的数据")
            
        return {"success": True, "data": step.step_data}
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ========== 短故事创作助手 API ==========

class ShortStorySettingsRequest(BaseModel):
    genre: str
    perspective: str = "first"
    summary: str = ""
    targetWords: int = 22000
    chapterCount: int = 8
    tropes: List[str] = []
    manuscript_id: Optional[int] = None


class ShortStoryOutlineRequest(BaseModel):
    settings: Dict[str, Any]


class ShortStoryChaptersRequest(BaseModel):
    settings: Dict[str, Any]
    outline: Dict[str, Any]



class ShortStoryReviewRequest(BaseModel):
    title: str
    intro: str
    chapters: List[Dict[str, Any]]
    settings: Dict[str, Any]
    project_id: Optional[int] = None
    manuscript_id: Optional[int] = None


class ShortStoryRewriteRequest(BaseModel):
    story_data: Dict[str, Any]
    review_report: str
    instruction: Optional[str] = None
    manuscript_id: Optional[int] = None  # 支持基于现有稿件重试


class ShortStoryStepRequest(BaseModel):
    manuscript_id: int
    step_name: str


@app.post("/api/short-story/generate-settings")
async def generate_short_story_settings(request: ShortStorySettingsRequest):
    """短故事: 生成设定 (含30字标题、极致人设)"""
    try:
        tropes_str = "、".join(request.tropes) if request.tropes else "自由发挥"
        perspective_str = "第一人称" if request.perspective == "first" else "第三人称"
        
        prompt = f"""你是专业的短故事创作专家，擅长创作新媒体爆款、反转不断的钩子文。请根据以下要求生成一个短故事设定。

**核心规范:**
- 字数限制: 严格{request.targetWords}字以内，一篇完结
- 节奏: 极速推进，无废话
- 视角: {perspective_str}
- 题材: {request.genre}
- 爆点梗: {tropes_str}

**用户灵感:**
{request.summary if request.summary else "由AI自由发挥创意"}

**新媒体爆款(Viral DNA)要求:**

1. **标题 (重中之重):**
   - 25字以内。必须极具冲击力。
   - 结构：[身份反差] + [极端反转/行为] + [悬念诱饵]
   - 示例："发现丈夫和小妹的聊天记录，我反手给他们办了婚礼，三年后他跪求复婚"
   - 严禁平淡。

2. **核心冲突:**
   - 必须包含“信息差” (只有主角或读者知道，反派不知道的真相)。
   - 包含“阶级/地位反转”或“情感背叛后的反杀”。

3. **人设极致:**
   - 主角：不能圣母，必须有强烈的反击意志或独特的复仇手段。
   - 反派：极度招人恨，其下场必须极其凄凉以满足“爽点”。

4. **文风:**
   - 黄金开篇：100字内必须爆发第一个矛盾点。
   - 短句为主，情绪饱满。

请严格按照以下JSON格式返回:
```json
{{
  "title": "爆款标题",
  "summary": "包含钩子和反转的情节简介",
  "main_conflict": "核心矛盾与信息差设计",
  "genre": "{request.genre}",
  "perspective": "{perspective_str}",
  "target_words": {request.targetWords},
  "chapter_count": {request.chapterCount},
  "tropes": {json.dumps(request.tropes, ensure_ascii=False)},
  "viral_dna": "本篇的爆款核心在哪（如：反转、打脸、悬疑）",
  "characters": [
    {{
      "name": "角色名",
      "role_type": "protagonist/antagonist/supporting",
      "identity": "反转前/后的双重身份",
      "personality": "极致性格描述",
      "flaw": "致命缺陷/伪装弱点"
    }}
  ],
  "golden_chapters": {{
    "chapter1": "第一章极强钩子",
    "chapter2": "第二章矛盾激化",
    "chapter3": "第三章反转/高潮"
  }}
}}
```
"""
        
        messages = [{"role": "user", "content": prompt}]
        response = ai_client._call_api(messages, temperature=0.9, max_tokens=4000)
        
        result = parse_json_response(response)
        
        if result:
            # === 保存稿件和步骤数据 ===
            db = next(get_db())
            
            # 如果有 manuscript_id，说明是重试，更新现有稿件
            if request.manuscript_id:
                manuscript = db.query(Manuscript).filter(Manuscript.id == request.manuscript_id).first()
                if manuscript:
                    manuscript.title = result.get('title', '未命名')
                    manuscript.updated_at = datetime.now()
            else:
                # 创建新稿件
                manuscript = Manuscript(
                    project_id=0, # 暂时为0，最后成文才关联具体项目
                    title=result.get('title', '未命名'),
                    status="generating"
                )
                db.add(manuscript)
                db.commit() # 获取ID
                db.refresh(manuscript)
            
            # 保存/更新 步骤数据
            step = ManuscriptStep(
                manuscript_id=manuscript.id,
                step_name="settings",
                step_data=result
            )
            db.add(step)
            db.commit()
            
            # 返回数据中带上 ID
            result['manuscript_id'] = manuscript.id
            return {"success": True, "data": result}
        else:
            return {"success": False, "message": "AI响应解析失败", "raw": response[:1000]}
            
    except HTTPException as e:
        raise e
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"生成设定失败: {str(e)}")


class ShortStoryOutlineRequest(BaseModel):
    settings: Dict[str, Any]
    manuscript_id: Optional[int] = None  # 新增

@app.post("/api/short-story/generate-outline")
async def generate_short_story_outline(request: ShortStoryOutlineRequest):
    """短故事: 生成大纲 (紧凑节奏，每章设钩子)"""
    try:
        settings = request.settings
        
        # 强制转换为整数，处理AI可能返回字符串的情况
        try:
            chapter_count = int(settings.get('chapter_count', 8))
            target_words = int(settings.get('target_words', 22000))
        except (ValueError, TypeError):
            chapter_count = 8
            target_words = 22000
            
        # 计算每章字数
        words_per_chapter = target_words // chapter_count
        
        prompt = f"""你是顶级网文架构师，擅长设计令人废寝忘食的爆款大纲。基于以下设定生成紧凑的章节大纲。

**小说设定:**
标题: {settings.get('title', '未命名')}
故事简介: {settings.get('summary', '')}
爆款核心: {settings.get('viral_dna', '')}
主要矛盾与信息差: {settings.get('main_conflict', '')}
题材: {settings.get('genre', '')}

**角色设定:**
{json.dumps(settings.get('characters', []), ensure_ascii=False, indent=2)}

**黄金三章设计:**
{json.dumps(settings.get('golden_chapters', {}), ensure_ascii=False, indent=2)}

**大纲架构要求:**
1. 共 {chapter_count} 章，总字数 {target_words} 字。
2. **钩子为王**: 每章结尾必须是一个强力悬念（钩子），让读者愿意“付费预览”或立即点击下一章。
3. **冲突高频**: 拒绝平淡描述，每章必须有一个完整的小冲突或大冲突的递进。
4. **信息差挖掘**: 利用角色间的信息不对称制造误会、反转或期待感。
5. **情绪曲线**: 情绪从“压抑/受辱”到“爆发/反击”的曲线要清晰。

请按JSON格式返回:
```json
{{
  "total_chapters": {chapter_count},
  "target_words": {target_words},
  "emotion_curve": [
    {{"chapter": 1, "intensity": 7, "type": "紧张/压抑"}}
  ],
  "chapters": [
    {{
      "chapter_number": 1,
      "title": "爆点标题",
      "summary": "包含核心冲突和行动的摘要",
      "target_words": {words_per_chapter},
      "hook": "本章强力悬念/反转点",
      "secondary_conflict": "本章制造的麻烦/阻碍",
      "emotion_intensity": 7,
      "emotion_type": "压抑/反击/爽点"
    }}
  ]
}}
```
"""
        
        messages = [{"role": "user", "content": prompt}]
        response = ai_client._call_api(messages, temperature=0.85, max_tokens=8000, timeout=600)
        
        # 尝试清理和解析
        result = parse_json_response(response)
        
        if result and result.get('chapters'):
            # === 保存步骤数据 ===
            if request.manuscript_id:
                try:
                    db = next(get_db())
                    step = ManuscriptStep(
                        manuscript_id=request.manuscript_id,
                        step_name="outline",
                        step_data=result
                    )
                    db.add(step)
                    db.commit()
                except Exception as db_e:
                    print(f"保存大纲步骤失败: {db_e}")

            return {"success": True, "data": result}
        else:

            # 如果解析失败，返回失败，让前端处理
            print(f"大纲JSON解析失败。原始响应长度: {len(response)}")
            return {
                "success": False, 
                "message": "大纲生成异常或格式错误，请尝试重新生成方案或修改摘要。",
                "raw": response[:1000]
            }
            
    except HTTPException as e:
        raise e
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"生成大纲失败: {str(e)}")


class ShortStoryChaptersRequest(BaseModel):
    settings: Dict[str, Any]
    outline: Dict[str, Any]
    manuscript_id: Optional[int] = None # 新增


@app.post("/api/short-story/generate-chapters")
async def generate_short_story_chapters(request: ShortStoryChaptersRequest):
    """短故事: 并行生成章节 (开头直入主题，心理描写重点)"""
    try:
        settings = request.settings
        outline = request.outline
        chapters = outline.get('chapters', [])
        
        if not chapters or not isinstance(chapters, list):
            raise HTTPException(status_code=400, detail="大纲中没有有效的章节信息")
        
        # 构建角色上下文
        characters = settings.get('characters', [])
        if not characters or not isinstance(characters, list):
            # 尝试修复或使用默认
            characters = []
            
        character_context = "\n".join([
            f"**{c.get('name', '未知角色')}**: {c.get('identity', '')} - {c.get('personality', '')}"
            for c in characters if isinstance(c, dict)
        ])
        
        perspective = settings.get('perspective', '第一人称')
        
        def generate_single_chapter(idx, chapter_outline):
            chapter_num = chapter_outline.get('chapter_number', idx + 1)
            title = chapter_outline.get('title', f'第{chapter_num}章')
            summary = chapter_outline.get('summary', '')
            target_words = chapter_outline.get('target_words', 2000)
            hook = chapter_outline.get('hook', '')
            
            is_first_chapter = (chapter_num == 1)
            
            prompt = f"""你是顶级网文写手，擅长新媒体爆款创作，文风犀利、情绪饱满、节奏感极强。
请创作第 {chapter_num} 章内容。

**小说核心:**
标题: {settings.get('title', '')}
爆款DNA: {settings.get('viral_dna', '')}
视角: {perspective}
本章标题: {title}

**本章任务:**
摘要: {summary}
必须包含的强力钩子/反转: {hook}
次要矛盾/冲突点: {chapter_outline.get('secondary_conflict', '')}

**人物关系图谱:**
{character_context}

**创作准则 (Viral Standard):**
1. **黄金开篇**: 每一章开头必须承接上文悬念，或立即开启新的小高潮。
2. **极速推进**: 严禁大段背景描写和心理独白，将所有情绪融入对话和动作中。
3. **情绪张力**: 每一段对话都要有潜台词，体现身份反差、压抑或爽点。
4. **视觉化动词**: 使用高饱和度的词汇，描述角色的微表情和肢体冲突。
5. **钩子收尾**: 本章结尾必须停在最扣人心弦的瞬间。

直接输出正文，不要有任何说明文字。
"""
            try:
                messages = [{"role": "user", "content": prompt}]
                content = ai_client._call_api(messages, temperature=0.85, max_tokens=4000)
                
                return {
                    "chapter_number": chapter_num,
                    "title": title,
                    "summary": summary,
                    "content": content,
                    "word_count": len(content),
                    "index": idx
                }
            except Exception as e:
                print(f"章节 {chapter_num} 生成失败: {e}")
                return None
        
        # 并行生成
        print(f"开始并行生成 {len(chapters)} 章短故事...")
        generated_chapters = []
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            future_to_idx = {
                executor.submit(generate_single_chapter, idx, ch): idx
                for idx, ch in enumerate(chapters)
            }
            
            results = []
            for future in concurrent.futures.as_completed(future_to_idx):
                result = future.result()
                if result:
                    results.append(result)
            
            results.sort(key=lambda x: x["index"])
            
            for res in results:
                del res["index"]
                generated_chapters.append(res)
        
        print(f"短故事章节生成完成，成功 {len(generated_chapters)}/{len(chapters)} 章")
        
        print(f"短故事章节生成完成，成功 {len(generated_chapters)}/{len(chapters)} 章")
        
        result_data = {"chapters": generated_chapters}

        # === 保存步骤数据 ===
        if request.manuscript_id:
            try:
                db = next(get_db())
                step = ManuscriptStep(
                    manuscript_id=request.manuscript_id,
                    step_name="chapters",
                    step_data=result_data
                )
                db.add(step)
                db.commit()
            except Exception as db_e:
                print(f"保存章节步骤失败: {db_e}")

        return {"success": True, "data": result_data}
        
    except HTTPException as e:
        raise e
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"生成章节失败: {str(e)}")


class ShortStoryNovelRequest(BaseModel):
    settings: Dict[str, Any]
    outline: Dict[str, Any]
    chapters: Dict[str, Any]
    manuscript_id: Optional[int] = None # 新增


@app.post("/api/short-story/generate-novel")
async def generate_short_story_novel(request: ShortStoryNovelRequest):
    """短故事: 一键成文 (创建项目并保存)"""
    try:
        settings = request.settings
        outline = request.outline
        chapters_data = request.chapters
        
        db = next(get_db())
        
        # 创建项目
        project = NovelProject(
            name=settings.get('title', '未命名短故事'),
            theme=settings.get('summary', ''),
            background=settings.get('genre', ''),
            genre="短故事",
            core_conflict=settings.get('main_conflict', ''),
            target_words=settings.get('target_words', 22000),
            status="completed",
            outline=outline,
            characters=settings.get('characters', []),
            chapters=[],
            word_count=0
        )
        
        db.add(project)
        db.commit()
        db.refresh(project)
        
        # 转换章节格式
        chapters = []
        total_words = 0
        
        for ch in chapters_data.get('chapters', []):
            chapter = {
                "id": ch.get('chapter_number', 1),
                "title": ch.get('title', ''),
                "summary": ch.get('summary', ''),
                "content": ch.get('content', ''),
                "word_count": ch.get('word_count', 0),
                "order": ch.get('chapter_number', 1)
            }
            chapters.append(chapter)
            total_words += ch.get('word_count', 0)
        
        project.chapters = chapters
        project.word_count = total_words
        db.commit()
        
        # 创建角色和章节数据
        for char in settings.get('characters', []):
            try:
                character = Character(
                    project_id=project.id,
                    name=char.get('name', '未命名'),
                    role_type=char.get('role_type', 'supporting'),
                    importance='core' if char.get('role_type') == 'protagonist' else 'important',
                    core_identity=char.get('identity'),
                    core_personality=char.get('personality'),
                    personality_flaw=char.get('flaw'),
                    source="short_story_assistant"
                )
                db.add(character)
            except Exception as e:
                print(f"创建角色失败: {e}")
        
        for idx, ch in enumerate(chapters_data.get('chapters', [])):
            try:
                plot_outline = PlotOutline(
                    project_id=project.id,
                    level="chapter",
                    chapter_number=ch.get('chapter_number', idx + 1),
                    title=ch.get('title', ''),
                    summary=ch.get('summary', ''),
                    target_words=ch.get('word_count', 2000),
                    source="short_story_assistant",
                    status="generated",
                    order=ch.get('chapter_number', idx + 1)
                )
                db.add(plot_outline)
                db.flush()
                
                chapter_draft = ChapterDraft(
                    project_id=project.id,
                    outline_id=plot_outline.id,
                    chapter_number=ch.get('chapter_number', idx + 1),
                    title=ch.get('title', ''),
                    content=ch.get('content', ''),
                    word_count=ch.get('word_count', 0),
                    status="completed",
                    generation_params={"source": "short_story_assistant"}
                )
                db.add(chapter_draft)
            except Exception as e:
                print(f"创建章节数据失败: {e}")
        
        db.commit()
        
        # 保存到稿件表 (Manuscript)
        try:
            manuscript = Manuscript(
                project_id=project.id,
                title=project.name,
                content=chapters,
                created_at=datetime.now(),
                updated_at=datetime.now()
            )
            db.add(manuscript)
            db.commit()
            db.refresh(manuscript)
            manuscript_id = manuscript.id
        except Exception as e:
            print(f"保存稿件记录失败: {e}")
            manuscript_id = None
        
        # 生成精彩导语
        try:
            intro_prompt = f"""请为这部短故事写一段精彩的推荐导语（Blurb）。
标题: {project.name}
题材: {project.background}
核心矛盾: {project.core_conflict}
简介: {project.theme}

要求:
1. 100-150字
2. 提炼故事最吸引人的亮点
3. 语言极具感染力，让人看一眼就想读下去
4. 也就是"最值得读"的理由

直接输出导语内容。"""
            
            intro_messages = [{"role": "user", "content": intro_prompt}]
            intro_content = ai_client._call_api(intro_messages, temperature=0.9, max_tokens=500)
            project.intro = intro_content # 假设模型有这个字段，或者暂时只返回给前端
        except Exception as e:
            print(f"导语生成失败: {e}")
            intro_content = project.theme

        return {
            "success": True,
            "data": {
                "project_id": project.id,
                "manuscript_id": manuscript_id,
                "title": project.name,
                "intro": intro_content,
                "chapters": chapters,
                "total_words": total_words
            },
            "message": "短故事生成完成!"
        }
        
    except HTTPException as e:
        raise e
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"整合小说失败: {str(e)}")


@app.post("/api/short-story/review")
async def review_short_story(request: ShortStoryReviewRequest):
    """短故事: AI 审稿 (结构-情绪 双轨版)"""
    try:
        chapters_text = "\n\n".join([
            f"### 第{c.get('chapter_number', i+1)}章: {c.get('title', '')}\n{c.get('content', '')}"
            for i, c in enumerate(request.chapters)
        ])
        
        prompt = f"""你是一位毒舌但极具实操经验的网文大主编，擅长从“结构”和“情绪”双线诊断爆款潜力。
请对以下短篇小说进行深度审稿。

**小说信息:**
标题: {request.title}
题材: {request.settings.get('genre', '')}
简介: {request.intro}

**正文内容:**
{chapters_text}

---

**审稿要求:**
你需要输出一份极其专业的审稿报告，格式必须严格参照以下模板：

1. **核心评级**: (S/A/B/C) + 核心短评
2. **第一部分：逐章结构与情绪双轨诊断**:
   - 对标文：说明该题材的黄金结构要求
   - 用户文诊断：逐章分析结构定位、情绪曲线、优缺点
3. **第二部分：四维核心问题诊断**:
   - 节奏与铺垫 (Pacing & Setup)
   - 人设与一致性 (Character & Consistency)
   - 冲突与爽点 (Conflict & Stakes)
   - 付费卡点 (Pay-for-Click Point)
4. **第三部分：核心修改方案**:
   - 明确指出哪些章节需要压缩，哪些高能剧情需要前置。

**注意:** 语言要犀利、专业、充满行业黑话（如：卡点、钩子、黄金三章、情绪断裂、种田陷阱等）。

请以 Markdown 格式输出。
"""
        messages = [{"role": "user", "content": prompt}]
        report = ai_client._call_api(messages, temperature=0.8, max_tokens=6000)
        
        # 提取评级
        grade = "B"
        if "**S**" in report or "评级：S" in report: grade = "S"
        elif "**A**" in report or "评级：A" in report: grade = "A"
        elif "**C**" in report or "评级：C" in report: grade = "C"
        
        # 如果提供了manuscript_id或project_id，保存审稿报告
        try:
            db = next(get_db())
            if request.manuscript_id:
                manuscript = db.query(Manuscript).filter(Manuscript.id == request.manuscript_id).first()
                if manuscript:
                    manuscript.review_report = report
                    manuscript.grade = grade
                    manuscript.updated_at = datetime.now()
                    db.commit()
            elif request.project_id:
                # 如果没有manuscript_id但有project_id，找到最新的一个稿件并更新
                manuscript = db.query(Manuscript).filter(Manuscript.project_id == request.project_id).order_by(Manuscript.created_at.desc()).first()
                if manuscript:
                    manuscript.review_report = report
                    manuscript.grade = grade
                    manuscript.updated_at = datetime.now()
                    db.commit()
        except Exception as db_err:
            print(f"保存审稿报告失败: {db_err}")

        return {"success": True, "data": {"report": report, "grade": grade}}
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"审稿失败: {str(e)}")


@app.get("/api/manuscripts/{project_id}")
async def list_manuscripts(project_id: int):
    """列出项目的所有稿件版本"""
    try:
        db = next(get_db())
        manuscripts = db.query(Manuscript).filter(Manuscript.project_id == project_id).order_by(Manuscript.created_at.desc()).all()
        return {
            "success": True, 
            "manuscripts": [
                {
                    "id": m.id,
                    "title": m.title,
                    "grade": m.grade,
                    "created_at": m.created_at.strftime("%Y-%m-%d %H:%M:%S")
                } for m in manuscripts
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/manuscript/{manuscript_id}")
async def get_manuscript(manuscript_id: int):
    """获取指定稿件的详细内容和审稿报告"""
    try:
        db = next(get_db())
        m = db.query(Manuscript).filter(Manuscript.id == manuscript_id).first()
        if not m:
            raise HTTPException(status_code=404, detail="稿件未找到")
        return {
            "success": True,
            "manuscript": {
                "id": m.id,
                "project_id": m.project_id,
                "title": m.title,
                "content": m.content,
                "review_report": m.review_report,
                "grade": m.grade,
                "created_at": m.created_at.strftime("%Y-%m-%d %H:%M:%S")
            }
        }
    except Exception as e:
        if isinstance(e, HTTPException): raise e
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/short-story/rewrite")
async def rewrite_short_story(request: ShortStoryRewriteRequest):
    """短故事: 根据审稿报告重写"""
    try:
        story = request.story_data
        report = request.review_report
        instruction = request.instruction or ""
        
        chapters = story.get('chapters', [])
        chapters_summary = "\n".join([
            f"- 第{c.get('chapter_number', i+1)}章: {c.get('title', '')} (约{c.get('word_count', 0)}字)"
            for i, c in enumerate(chapters)
        ])
        
        # 1. 生成重写计划 (JSON)
        plan_prompt = f"""你是顶级网文主编。请根据“审稿报告”，为这部短篇小说的每一章制定详细的“重写指令”。
        
**审稿报告建议:**
{report}

**额外修改要求:**
{instruction}

**原稿章节列表:**
{chapters_summary}

**任务:**
请输出一份JSON格式的重写计划，明确每一章需要如何根据建议进行具体调整（如：将冲突从第三章移到第一章，强化主角在第二章的反击等）。

格式如下：
```json
{{
  "rewrite_plan": [
    {{
      "chapter_number": 1,
      "original_title": "旧标题",
      "new_title": "新爆款标题",
      "modifications": "具体修改点：1... 2... 3...",
      "target_hook": "本章必须强化的悬念/钩子"
    }}
  ]
}}
```
"""
        messages = [{"role": "user", "content": plan_prompt}]
        plan_response = ai_client._call_api(messages, temperature=0.7, max_tokens=2000)
        plan_result = parse_json_response(plan_response)
        
        if not plan_result or not plan_result.get('rewrite_plan'):
            # 记录原始响应以便调试
            with open("plan_response_error.log", "w") as f:
                f.write(f"RAW RESPONSE:\n{plan_response}\n\nPARSED RESULT:\n{json.dumps(plan_result, ensure_ascii=False)}")
            raise HTTPException(status_code=500, detail="生成重写计划失败，已记录原始响应")

        rewrite_plans = plan_result['rewrite_plan']
        
        # 2. 并行重写每一章
        def execute_chapter_rewrite(idx, plan):
            orig_chapter = chapters[idx] if idx < len(chapters) else {}
            chapter_num = plan.get('chapter_number', idx + 1)
            
            chapter_prompt = f"""你是顶级网文修改专家。请执行具体的章节重写。
            
**原章节内容 (第{chapter_num}章):**
标题: {orig_chapter.get('title', '无')}
正文:
{orig_chapter.get('content', '无内容')}

**重写计划要求:**
新标题: {plan.get('new_title', '爆款标题')}
具体修改指令: {plan.get('modifications', '全方位升级')}
核心钩子: {plan.get('target_hook', '留悬念')}

**重写规则:**
1. 保持视角一致。
2. 节奏极速，砍掉废话。
3. 心理描写融入行动。
4. 结尾必须卡在钩子上。

直接输出重写后的章节正文，不要有任何说明。
"""
            try:
                msg = [{"role": "user", "content": chapter_prompt}]
                new_content = ai_client._call_api(msg, temperature=0.85, max_tokens=4000)
                return {
                    "chapter_number": chapter_num,
                    "title": plan.get('new_title', '新标题'),
                    "content": new_content,
                    "word_count": len(new_content)
                }
            except Exception as e:
                print(f"重写章节 {chapter_num} 失败: {e}")
                return None

        # 并行执行
        new_chapters = []
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            futures = [
                executor.submit(execute_chapter_rewrite, i, plan)
                for i, plan in enumerate(rewrite_plans)
            ]
            
            for future in concurrent.futures.as_completed(futures):
                res = future.result()
                if res:
                    new_chapters.append(res)
        
        # 排序
        new_chapters.sort(key=lambda x: x['chapter_number'])
        
        if new_chapters:
            return {"success": True, "data": {"chapters": new_chapters}}
        else:
            return {"success": False, "message": "所有章节重写均失败"}
            
    except Exception as e:
        import traceback
        error_msg = traceback.format_exc()
        print(error_msg)
        with open("rewrite_debug.log", "w") as f:
            f.write(error_msg)
        raise HTTPException(status_code=500, detail=f"重构失败: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host=settings.host, port=settings.port)


# ========== 角色卡系统（星月风格）==========

class CharacterCardCreate(BaseModel):
    """创建角色卡"""
    project_id: int
    name: str
    role_type: str
    importance: str = "supporting"
    core_identity: Optional[str] = None
    core_personality: Optional[str] = None
    core_motivation: Optional[str] = None
    personality_flaw: Optional[str] = None
    flaw_consequence: Optional[str] = None
    age: Optional[int] = None
    gender: Optional[str] = None


@app.post("/api/character-cards")
async def create_character_card(request: CharacterCardCreate):
    """创建角色卡"""
    try:
        db = next(get_db())
        character = Character(
            project_id=request.project_id,
            name=request.name,
            role_type=request.role_type,
            importance=request.importance,
            core_identity=request.core_identity,
            core_personality=request.core_personality,
            core_motivation=request.core_motivation,
            personality_flaw=request.personality_flaw,
            flaw_consequence=request.flaw_consequence,
            age=request.age,
            gender=request.gender,
            source="manual"
        )
        db.add(character)
        db.commit()
        return {"success": True, "data": {"id": character.id, "name": character.name}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/character-cards/{project_id}")
async def get_character_cards(project_id: int):
    """获取角色卡列表（星月风格：智能过滤 + 按重要性排序）"""
    try:
        db = next(get_db())

        # 使用 CASE 语句实现正确的排序：core → important → supporting
        from sqlalchemy import case

        importance_order = case(
            (Character.importance == 'core', 1),
            (Character.importance == 'important', 2),
            (Character.importance == 'supporting', 3),
            else_=4
        )

        characters = db.query(Character).filter(
            Character.project_id == project_id,
            Character.is_visible == 1
        ).order_by(importance_order, Character.id).all()

        result = []
        for char in characters:
            result.append({
                "id": char.id,
                "name": char.name,
                "role_type": char.role_type,
                "importance": char.importance,
                "status": char.status,
                "core_identity": char.core_identity,
                "core_personality": char.core_personality,
                "core_motivation": char.core_motivation,
                "personality_flaw": char.personality_flaw,
                "flaw_consequence": char.flaw_consequence,
                "age": char.age,
                "gender": char.gender,
                "current_location": char.current_location,
                "growth_direction": char.growth_direction,
                "speech_example": char.speech_example,
                "relationship_notes": char.relationship_notes,
                "source": char.source
            })
        return {"success": True, "data": result}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/character-cards/{character_id}")
async def update_character_card(character_id: int, request: dict):
    """更新角色卡"""
    try:
        db = next(get_db())
        character = db.query(Character).filter(Character.id == character_id).first()
        if not character:
            raise HTTPException(status_code=404, detail="角色卡不存在")
        
        for key, value in request.items():
            if hasattr(character, key) and value is not None:
                setattr(character, key, value)
        
        db.commit()
        return {"success": True, "message": "更新成功"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/character-cards/{character_id}")
async def delete_character_card(character_id: int):
    """删除角色卡"""
    try:
        db = next(get_db())
        character = db.query(Character).filter(Character.id == character_id).first()
        if not character:
            raise HTTPException(status_code=404, detail="角色卡不存在")
        db.delete(character)
        db.commit()
        return {"success": True, "message": "删除成功"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# =====================================================================
# 渠道投稿智能体 API (Channel Agent)
# =====================================================================

class ChannelAgentCreate(BaseModel):
    """创建渠道智能体"""
    name: str
    description: Optional[str] = None
    channel_type: Optional[str] = None
    target_audience: Optional[str] = None
    channel_characteristics: Optional[Dict[str, Any]] = None
    temperature: Optional[int] = 70
    top_p: Optional[int] = 90


class ChannelAgentUpdate(BaseModel):
    """更新渠道智能体"""
    name: Optional[str] = None
    description: Optional[str] = None
    channel_type: Optional[str] = None
    target_audience: Optional[str] = None
    channel_characteristics: Optional[Dict[str, Any]] = None
    temperature: Optional[int] = None
    top_p: Optional[int] = None
    is_active: Optional[int] = None


@app.post("/api/channel-agents")
async def create_channel_agent(agent: ChannelAgentCreate):
    """创建渠道智能体"""
    try:
        db = next(get_db())
        new_agent = ChannelAgent(
            name=agent.name,
            description=agent.description,
            channel_type=agent.channel_type,
            target_audience=agent.target_audience,
            channel_characteristics=agent.channel_characteristics,
            temperature=agent.temperature,
            top_p=agent.top_p,
            training_status="pending"
        )
        db.add(new_agent)
        db.commit()
        db.refresh(new_agent)

        return {
            "success": True,
            "data": {
                "id": new_agent.id,
                "name": new_agent.name,
                "training_status": new_agent.training_status
            }
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/channel-agents")
async def list_channel_agents():
    """获取所有渠道智能体列表"""
    try:
        db = next(get_db())
        agents = db.query(ChannelAgent).filter(
            ChannelAgent.is_active == 1
        ).order_by(ChannelAgent.id.desc()).all()

        result = []
        for agent in agents:
            result.append({
                "id": agent.id,
                "name": agent.name,
                "description": agent.description,
                "channel_type": agent.channel_type,
                "target_audience": agent.target_audience,
                "channel_characteristics": agent.channel_characteristics,
                "length_requirements": agent.length_requirements,
                "contact_info": agent.contact_info,
                "training_status": agent.training_status,
                "corpus_word_count": agent.corpus_word_count,
                "usage_count": agent.usage_count,
                "success_count": agent.success_count,
                "created_at": agent.created_at.isoformat() if agent.created_at else None
            })

        return {"success": True, "data": result}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/channel-agents/{agent_id}")
async def get_channel_agent(agent_id: int):
    """获取渠道智能体详情"""
    try:
        db = next(get_db())
        agent = db.query(ChannelAgent).filter(ChannelAgent.id == agent_id).first()
        if not agent:
            raise HTTPException(status_code=404, detail="智能体不存在")

        return {"success": True, "data": agent}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/channel-agents/{agent_id}/upload-corpus")
async def upload_corpus(agent_id: int, files: List[UploadFile] = File(...)):
    """上传语料文件（支持多文件）"""
    try:
        db = next(get_db())
        agent = db.query(ChannelAgent).filter(ChannelAgent.id == agent_id).first()
        if not agent:
            raise HTTPException(status_code=404, detail="智能体不存在")

        upload_dir = "uploads/corpus"
        os.makedirs(upload_dir, exist_ok=True)

        if not agent.training_files:
            agent.training_files = []

        total_word_count = 0
        uploaded_files = []

        for file in files:
            file_ext = file.filename.split('.')[-1]
            safe_filename = f"agent_{agent_id}_{int(datetime.now().timestamp())}_{file.filename}"
            file_path = os.path.join(upload_dir, safe_filename)

            with open(file_path, "wb") as f:
                content = await file.read()
                f.write(content)

            word_count = len(content)
            total_word_count += word_count

            agent.training_files.append({
                "filename": safe_filename,
                "original_name": file.filename,
                "upload_time": datetime.now().isoformat(),
                "word_count": word_count
            })

            uploaded_files.append({
                "filename": safe_filename,
                "original_name": file.filename,
                "word_count": word_count
            })

        agent.corpus_word_count = (agent.corpus_word_count or 0) + total_word_count
        db.commit()

        return {
            "success": True,
            "message": f"成功上传{len(files)}个语料文件",
            "data": {
                "uploaded_count": len(files),
                "total_word_count": total_word_count,
                "files": uploaded_files
            }
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ========== 长文生成相关 ==========

from backend.generator.long_novel_generator import LongNovelGenerator
long_novel_generator = LongNovelGenerator()

class CreateLongProjectRequest(BaseModel):
    manuscript_id: int
    title: str

class ExpandVolumeRequest(BaseModel):
    project_id: int
    chapter_index: int # 短篇章节索引 (0-7)

class GenerateLongChapterRequest(BaseModel):
    project_id: int
    chapter_id: int # 数据库ID

# ========== 长篇扩写助手 (Long Story Assistant) ==========

expansion_engine = ExpansionEngine()

@app.post("/api/long-novel/create")
async def create_long_project(request: CreateLongProjectRequest):
    """从短篇稿件创建长篇项目"""
    db_gen = get_db()
    db = next(db_gen)
    try:
        manuscript = db.query(Manuscript).filter(Manuscript.id == request.manuscript_id).first()
        if not manuscript:
            return {"success": False, "message": "找不到原始稿件"}
            
        # 创建长篇项目
        new_project = NovelProject(
            name=request.title,
            type="long_novel",
            source_manuscript_id=manuscript.id,
            status="planning",
            outline={"source": "short_story_expansion"},
            chapters=[],
            word_count=0
        )
        db.add(new_project)
        db.commit()
        db.refresh(new_project)
        
        return {"success": True, "data": {"project_id": new_project.id}}
    except Exception as e:
        return {"success": False, "message": str(e)}
    finally:
        db.close()

@app.get("/api/long-novel/{project_id}")
async def get_long_project(project_id: int):
    """获取长篇项目详情 (包含卷和章节结构)"""
    db_gen = get_db()
    db = next(db_gen)
    try:
        project = db.query(NovelProject).filter(NovelProject.id == project_id).first()
        if not project:
            return {"success": False, "message": "项目不存在"}

        # 获取映射关系 (卷结构)
        mappings = db.query(LongNovelMapping).filter(LongNovelMapping.project_id == project_id).order_by(LongNovelMapping.volume_number).all()
        
        # 获取所有生成的大纲/章节
        # 这里为了简化，我们假设 plot_outlines 表存了大纲，chapter_drafts 存了正文
        # 实际需要根据前面的数据库设计来关联
        
        # 构造返回结构
        volumes = []
        for m in mappings:
            # 查询本卷下的章节
            chapters = db.query(ChapterDraft).join(PlotOutline).filter(
                PlotOutline.project_id == project_id,
                ChapterDraft.chapter_number >= m.start_chapter,
                ChapterDraft.chapter_number <= m.end_chapter
            ).all()
            
            # 手动组装 chapters info 
            # (这里简化处理，实际应该 join 查询 PlotOutline 获取标题和摘要)
            chapter_list = []
            outlines = db.query(PlotOutline).filter(
                PlotOutline.project_id == project_id,
                PlotOutline.chapter_number >= m.start_chapter,
                PlotOutline.chapter_number <= m.end_chapter
            ).order_by(PlotOutline.chapter_number).all()
            
            for ol in outlines:
                draft = db.query(ChapterDraft).filter(ChapterDraft.outline_id == ol.id).first()
                focus = json.loads(ol.focus_elements) if ol.focus_elements else {}
                chapter_list.append({
                    "id": draft.id if draft else None,
                    "outline_id": ol.id,
                    "chapter_number": ol.chapter_number,
                    "title": ol.title,
                    "summary": ol.summary,
                    "main_conflict": focus.get("main_conflict"),
                    "sub_conflict": focus.get("sub_conflict"),
                    "emotion_arc": focus.get("emotion_arc"),
                    "status": draft.status if draft else "planning",
                    "word_count": draft.word_count if draft else 0,
                    "content": draft.content if draft else ""
                })
            
            volumes.append({
                "volume_number": m.volume_number,
                "title": m.short_chapter_title,
                "summary": m.short_chapter_summary,
                "chapters": chapter_list
            })
            
        return {
            "success": True, 
            "data": {
                "project": {
                    "id": project.id,
                    "name": project.name,
                    "source_manuscript_id": project.source_manuscript_id
                },
                "volumes": volumes
            }
        }
    finally:
        db.close()

@app.post("/api/long-novel/expand-volume")
async def expand_volume(request: ExpandVolumeRequest):
    """
    【新算法逻辑】将短篇的一章扩写为长篇的一卷 (裂变为 18-20 章)
    """
    db_gen = get_db()
    db = next(db_gen)
    try:
        project = db.query(NovelProject).filter(NovelProject.id == request.project_id).first()
        if not project or not project.source_manuscript_id:
            return {"success": False, "message": "无效的长篇项目"}
            
        manuscript = db.query(Manuscript).filter(Manuscript.id == project.source_manuscript_id).first()
        if not manuscript:
            return {"success": False, "message": "找不到源短篇稿件"}
            
        # 1. 使用 ExpansionEngine 裂变大纲 (1:18-20)
        try:
            expanded_data = expansion_engine.plan_volume_expansion(manuscript, request.chapter_index)
        except Exception as ai_err:
            return {"success": False, "message": f"Expansion planning failed: {str(ai_err)}"}
            
        if not expanded_data:
            return {"success": False, "message": "AI生成结果为空"}
            
        # 2. 计算章节号范围 (累积计算)
        last_mapping = db.query(LongNovelMapping).filter(LongNovelMapping.project_id == request.project_id).order_by(LongNovelMapping.end_chapter.desc()).first()
        start_chap = (last_mapping.end_chapter + 1) if last_mapping else 1
        
        chapters_data = expanded_data.get("chapters", [])
        num_new_chapters = len(chapters_data)
        end_chap = start_chap + num_new_chapters - 1
        
        # 3. 保存映射
        short_chapters = manuscript.content.get('chapters', [])
        short_chap_title = short_chapters[request.chapter_index].get('title') if request.chapter_index < len(short_chapters) else f"第{request.chapter_index+1}章"
        short_chap_summary = short_chapters[request.chapter_index].get('summary', "") if request.chapter_index < len(short_chapters) else ""
        
        mapping = LongNovelMapping(
            project_id=project.id,
            short_chapter_title=short_chap_title,
            short_chapter_summary=short_chap_summary,
            volume_number=request.chapter_index + 1,
            start_chapter=start_chap,
            end_chapter=end_chap
        )
        db.add(mapping)
        db.flush()
        
        # 4. 批量创建 PlotOutline 和 ChapterDraft
        for i, ch_data in enumerate(chapters_data):
            curr_ch_num = start_chap + i
            outline = PlotOutline(
                project_id=project.id,
                level="chapter",
                parent_id=mapping.id,
                chapter_number=curr_ch_num,
                title=ch_data.get("title"),
                summary=ch_data.get("summary"),
                target_words=3000,
                focus_elements=json.dumps({
                    "main_conflict": ch_data.get("main_conflict"),
                    "sub_conflict": ch_data.get("sub_conflict"),
                    "emotion_arc": ch_data.get("emotion_arc")
                }),
                status="ready"
            )
            db.add(outline)
            db.flush()
            
            draft = ChapterDraft(
                project_id=project.id,
                outline_id=outline.id,
                chapter_number=curr_ch_num,
                title=ch_data.get("title"),
                status="draft"
            )
            db.add(draft)
            
        db.commit()
        return {"success": True, "data": {"volume_number": mapping.volume_number, "num_chapters": num_new_chapters}}
        
    except Exception as e:
        db.rollback()
        import traceback
        traceback.print_exc()
        return {"success": False, "message": str(e)}
    finally:
        db.close()

@app.post("/api/long-novel/preview-outline")
async def preview_long_outline(request: CreateLongProjectRequest):
    """
    【预览功能】生成完整的长篇扩写大纲预览（不保存到数据库）

    用于在创建长篇项目前，让用户预览整体结构规划
    返回完整的180章规划结构，方便用户确认是否符合预期
    """
    db_gen = get_db()
    db = next(db_gen)
    try:
        manuscript = db.query(Manuscript).filter(Manuscript.id == request.manuscript_id).first()
        if not manuscript:
            return {"success": False, "message": "找不到原始稿件"}

        # 使用 ExpansionEngine 生成完整大纲规划
        outline_plan = expansion_engine.generate_full_outline_plan(manuscript)

        return {
            "success": True,
            "data": outline_plan
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"success": False, "message": str(e)}
    finally:
        db.close()


@app.post("/api/long-novel/generate-chapter")
async def generate_long_chapter(request: GenerateLongChapterRequest):
    """【新扩写策略】生成长篇单章正文 (3000字+)"""
    db_gen = get_db()
    db = next(db_gen)
    try:
        # 1. 获取草稿和关联大纲
        draft = db.query(ChapterDraft).filter(ChapterDraft.id == request.chapter_id).first()
        if not draft:
            return {"success": False, "message": "章节不存在"}
            
        outline = db.query(PlotOutline).filter(PlotOutline.id == draft.outline_id).first()
        project = db.query(NovelProject).filter(NovelProject.id == draft.project_id).first()
        
        # 2. 获取卷信息 (Volume mapping)
        mapping = db.query(LongNovelMapping).filter(
            LongNovelMapping.project_id == project.id,
            LongNovelMapping.start_chapter <= draft.chapter_number,
            LongNovelMapping.end_chapter >= draft.chapter_number
        ).first()
        
        volume_info = {
            "volume_title": mapping.short_chapter_title if mapping else "未知卷",
            "volume_summary": mapping.short_chapter_summary if mapping else ""
        }
        
        # 3. 构造上下文 (前3章正文)
        prev_drafts = db.query(ChapterDraft).filter(
            ChapterDraft.project_id == project.id,
            ChapterDraft.chapter_number < draft.chapter_number
        ).order_by(ChapterDraft.chapter_number.desc()).limit(3).all()
        
        context = "前情提要:\n" + "\n".join([f"第{d.chapter_number}章 {d.title}: {d.content[:500]}..." for d in reversed(prev_drafts)])
        
        # 4. 获取人物信息 (从原始稿件中提取)
        characters = []
        manuscript = db.query(Manuscript).filter(Manuscript.id == project.source_manuscript_id).first()
        if manuscript and manuscript.content:
            characters = manuscript.content.get('characters', [])
        
        # 5. 调用 ExpansionEngine 执行高质量扩写
        focus_elements = json.loads(outline.focus_elements) if outline.focus_elements else {}
        chapter_outline_info = {
            "title": outline.title,
            "summary": outline.summary,
            "main_conflict": focus_elements.get("main_conflict", ""),
            "sub_conflict": focus_elements.get("sub_conflict", ""),
            "emotion_arc": focus_elements.get("emotion_arc", "")
        }
        
        chapter_content = expansion_engine.generate_long_chapter(
            volume_info=volume_info,
            chapter_outline=chapter_outline_info,
            context=context,
            characters=characters
        )
        
        # 6. 保存正文
        draft.content = chapter_content
        draft.word_count = len(chapter_content)
        draft.status = "completed"
        draft.updated_at = datetime.now()
        
        db.commit()
        
        return {"success": True, "data": {"content": chapter_content, "word_count": len(chapter_content)}}
        
    except Exception as e:
        return {"success": False, "message": str(e)}
    finally:
        db.close()


@app.post("/api/channel-agents/{agent_id}/train")
async def train_channel_agent(agent_id: int):
    """AI拆解训练 - 从语料中提取风格特征"""
    try:
        db = next(get_db())
        agent = db.query(ChannelAgent).filter(ChannelAgent.id == agent_id).first()
        if not agent:
            raise HTTPException(status_code=404, detail="智能体不存在")

        if not agent.training_files:
            raise HTTPException(status_code=400, detail="请先上传语料文件")

        agent.training_status = "training"
        agent.training_progress = 0
        db.commit()

        async def train_agent():
            try:
                corpus_texts = []
                for file_info in agent.training_files:
                    file_path = os.path.join("uploads/corpus", file_info["filename"])
                    if os.path.exists(file_path):
                        with open(file_path, "r", encoding="utf-8") as f:
                            corpus_texts.append(f.read())

                full_corpus = "\n\n".join(corpus_texts)
                ai_client = DeepSeekClient(api_key=settings.deepseek_api_key, model=settings.deepseek_model)

                analysis_prompt = f"""分析以下公众号语料风格，以JSON格式返回特征：\n\n{full_corpus[:10000]}\n\n请提取：title_style, topic_preferences, writing_style, content_structure, length_requirements, vocabulary_features"""

                agent.training_progress = 50
                db.commit()

                response = ai_client.chat(analysis_prompt, temperature=0.3)
                
                import json
                json_match = re.search(r'\{.*\}', response, re.DOTALL)
                if json_match:
                    analysis_result = json.loads(json_match.group())
                    agent.title_style = analysis_result.get("title_style")
                    agent.topic_preferences = analysis_result.get("topic_preferences")
                    agent.writing_style = analysis_result.get("writing_style")
                    agent.content_structure = analysis_result.get("content_structure")
                    agent.length_requirements = analysis_result.get("length_requirements")
                    agent.vocabulary_features = analysis_result.get("vocabulary_features")

                agent.training_progress = 100
                agent.training_status = "completed"
                agent.last_training_at = datetime.now()
                db.commit()

            except Exception as e:
                agent.training_status = "failed"
                agent.training_error = str(e)
                db.commit()

        import asyncio
        asyncio.create_task(train_agent())

        return {"success": True, "message": "训练已开始"}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/channel-agents/{agent_id}/generate-inspiration")
async def generate_inspiration(agent_id: int, request: dict):
    """使用渠道智能体生成文章灵感/大纲"""
    try:
        db = next(get_db())
        agent = db.query(ChannelAgent).filter(ChannelAgent.id == agent_id).first()
        if not agent:
            raise HTTPException(status_code=404, detail="智能体不存在")

        if agent.training_status != "completed":
            raise HTTPException(status_code=400, detail="智能体尚未完成训练")

        topic = request.get("topic", "")  # 可选
        requirements = request.get("additional_requirements", "")

        # 提取智能体的风格特征
        style_info = agent.writing_style or {}
        title_style = agent.title_style or {}
        content_structure = agent.content_structure or {}
        topic_preferences = agent.topic_preferences or {}

        # 构建灵感生成的prompt
        if topic:
            # 如果提供了主题，基于主题生成灵感
            inspiration_prompt = f"""你是一位专业的公众号编辑，请根据以下信息生成文章创作灵感：

【主题】{topic}

【渠道风格特征】
- 渠道名称：{agent.name}
- 标题风格：{title_style.get('characteristics', [])}
- 标题示例：{title_style.get('examples', [])}
- 写作风格：{style_info.get('tone', '')}
- 叙述视角：{style_info.get('perspective', '')}
- 内容结构：{content_structure.get('type', '默认结构')}
- 段落密度：{content_structure.get('paragraph_density', '')}
- 受众定位：{agent.target_audience or '普通读者'}
- 核心主题偏好：{topic_preferences.get('core_themes', [])}
- 情绪关键词：{topic_preferences.get('emotional_keywords', [])}

【额外要求】
{requirements if requirements else '无'}

请生成一个详细的创作灵感，包括以下内容，以JSON格式返回：

{{
    "topic_suggestion": "{topic}",
    "title_suggestions": [
        "标题建议1（符合该渠道风格）",
        "标题建议2",
        "标题建议3"
    ],
    "core_angle": "核心角度和切入点（50-100字）",
    "content_outline": [
        {{"section": "第一部分标题", "key_points": "要点说明", "emotional_tone": "情绪基调"}},
        {{"section": "第二部分标题", "key_points": "要点说明", "emotional_tone": "情绪基调"}},
        {{"section": "第三部分标题", "key_points": "要点说明", "emotional_tone": "情绪基调"}}
    ],
    "key_elements": ["关键元素1", "关键元素2", "关键元素3"],
    "emotional_arc": "情绪弧线设计说明",
    "estimated_words": 预计字数,
    "writing_notes": "创作要点提醒"
}}

只返回JSON，不要有其他说明文字。"""
        else:
            # 没有提供主题，让AI自主生成主题和灵感
            inspiration_prompt = f"""你是一位专业的公众号编辑，请根据该渠道的风格特征，自主创作一个合适的文章主题和详细的创作灵感。

【渠道风格特征】
- 渠道名称：{agent.name}
- 渠道描述：{agent.description or ''}
- 标题风格：{title_style.get('characteristics', [])}
- 标题示例：{title_style.get('examples', [])}
- 写作风格：{style_info.get('tone', '')}
- 叙述视角：{style_info.get('perspective', '')}
- 语言特点：{style_info.get('language_features', {})}
- 内容结构：{content_structure.get('type', '默认结构')}
- 开头方式：{content_structure.get('opening', '')}
- 段落密度：{content_structure.get('paragraph_density', '')}
- 受众定位：{agent.target_audience or '普通读者'}
- 核心主题偏好：{topic_preferences.get('core_themes', [])}
- 情绪关键词：{topic_preferences.get('emotional_keywords', [])}
- 避免元素：{topic_preferences.get('avoid_elements', [])}

请自主创作一个符合该渠道定位的文章主题，并生成详细的创作灵感，以JSON格式返回：

{{
    "topic_suggestion": "AI自主推荐的文章主题（符合该渠道定位和受众）",
    "title_suggestions": [
        "标题建议1（完全符合该渠道标题风格）",
        "标题建议2",
        "标题建议3"
    ],
    "core_angle": "核心角度和切入点（50-100字，说明为什么这个主题适合这个渠道）",
    "content_outline": [
        {{"section": "第一部分标题", "key_points": "要点说明", "emotional_tone": "情绪基调"}},
        {{"section": "第二部分标题", "key_points": "要点说明", "emotional_tone": "情绪基调"}},
        {{"section": "第三部分标题", "key_points": "要点说明", "emotional_tone": "情绪基调"}}
    ],
    "key_elements": ["关键元素1", "关键元素2", "关键元素3"],
    "emotional_arc": "情绪弧线设计说明",
    "estimated_words": 预计字数,
    "writing_notes": "创作要点提醒"
}}

只返回JSON，不要有其他说明文字。"""

        ai_client = DeepSeekClient(api_key=settings.deepseek_api_key, model=settings.deepseek_model)
        response = ai_client.chat(inspiration_prompt, temperature=0.9, max_tokens=2500)  # 提高温度增加创意

        # 尝试解析JSON
        import re
        json_match = re.search(r'\{[\s\S]*\}', response)
        if json_match:
            try:
                import json
                inspiration_data = json.loads(json_match.group())
                return {"success": True, "data": inspiration_data}
            except json.JSONDecodeError:
                pass

        # 如果解析失败，返回原始响应
        return {"success": True, "data": {"raw_inspiration": response}}

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/channel-agents/{agent_id}/generate")
async def generate_with_agent(agent_id: int, request: dict):
    """使用渠道智能体生成投稿文章"""
    try:
        db = next(get_db())
        agent = db.query(ChannelAgent).filter(ChannelAgent.id == agent_id).first()
        if not agent:
            raise HTTPException(status_code=404, detail="智能体不存在")

        if agent.training_status != "completed":
            raise HTTPException(status_code=400, detail="智能体尚未完成训练")

        topic = request.get("topic", "")
        inspiration = request.get("inspiration", "")  # 可选的灵感
        word_count = request.get("word_count")

        # 如果没有主题也没有灵感，提示用户先生成灵感
        if not topic and not inspiration:
            raise HTTPException(status_code=400, detail="请先生成灵感或提供文章主题")

        # 构建生成prompt
        if inspiration:
            # 如果提供了灵感，基于灵感生成（灵感中已包含主题）
            topic_part = f'【文章主题】{topic}\n\n' if topic else ''
            prompt = f"""请根据以下详细的创作灵感生成完整的公众号文章：

{topic_part}【创作灵感】
{inspiration}

【渠道风格要求】
- 标题风格：{agent.title_style or '常规标题'}
- 写作风格：{agent.writing_style or '自然流畅'}
- 内容结构：{agent.content_structure or '常规结构'}
- 目标字数：{word_count or '根据内容自定'}

请直接生成文章，不要有任何说明和过渡，直接进入内容！"""
        else:
            # 没有灵感，直接根据主题生成
            prompt = f"""根据风格要求生成公众号文章：
主题：{topic}
风格：{agent.writing_style}
结构：{agent.content_structure}
字数：{word_count or agent.length_requirements}

生成完整文章。"""

        ai_client = DeepSeekClient(api_key=settings.deepseek_api_key, model=settings.deepseek_model)
        response = ai_client.chat(prompt, temperature=agent.temperature/100 if hasattr(agent, 'temperature') and agent.temperature else 0.7, max_tokens=4000)

        agent.usage_count = (agent.usage_count or 0) + 1
        db.commit()

        return {"success": True, "data": {"content": response, "word_count": len(response)}}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ========== 仿写生成系统 ==========

from backend.api.imitation_api import (
    ImitationGenerator,
    DeconstructionRequest,
    ConfigurationRequest,
    PreviewRequest,
    GenerationRequest
)


@app.post("/api/imitation/deconstruct")
async def deconstruct_original(request: DeconstructionRequest):
    """阶段一：深度拆解原文"""
    try:
        db = next(get_db())
        generator = ImitationGenerator(db)
        result = generator.deconstruct(request)

        if result.success:
            return {"success": True, "project_id": result.project_id, "analysis": result.analysis}
        else:
            return {"success": False, "message": result.message}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/imitation/configure")
async def configure_imitation(request: ConfigurationRequest):
    """阶段二：配置新设定"""
    try:
        db = next(get_db())
        generator = ImitationGenerator(db)
        result = generator.configure(request)

        if result.success:
            return {"success": True, "message": result.message}
        else:
            return {"success": False, "message": result.message}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/imitation/preview")
async def preview_reconstruction(request: PreviewRequest):
    """阶段三：生成重构蓝图预览"""
    try:
        db = next(get_db())
        generator = ImitationGenerator(db)
        result = generator.preview(request)

        if result.success:
            return {"success": True, "blueprint": result.blueprint}
        else:
            return {"success": False, "message": result.message}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/imitation/generate")
async def generate_imitation(request: GenerationRequest):
    """阶段四：生成仿写正文"""
    try:
        db = next(get_db())
        generator = ImitationGenerator(db)
        result = generator.generate(request)

        if result.success:
            return {"success": True, "content": result.content}
        else:
            return {"success": False, "message": result.message}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/imitation/projects/{project_id}")
async def get_imitation_project(project_id: int):
    """获取仿写项目详情"""
    try:
        db = next(get_db())
        project = db.query(ImitationProject).filter(ImitationProject.id == project_id).first()

        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        return {
            "success": True,
            "data": {
                "id": project.id,
                "title": project.title,
                "status": project.status,
                "original_title": project.original_title,
                "original_content": project.original_content,
                "new_worldview": project.new_worldview,
                "protagonist_setting": project.protagonist_setting,
                "core_conflict": project.core_conflict,
                "golden_finger": project.golden_finger,
                "deconstruction_result": project.deconstruction_result,
                "reconstruction_blueprint": project.reconstruction_blueprint,
                "generated_content": project.generated_content,
                "created_at": project.created_at.isoformat(),
                "updated_at": project.updated_at.isoformat()
            }
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/imitation/projects")
async def list_imitation_projects():
    """获取所有仿写项目列表"""
    try:
        db = next(get_db())
        projects = db.query(ImitationProject).order_by(ImitationProject.created_at.desc()).all()

        return {
            "success": True,
            "data": [
                {
                    "id": p.id,
                    "title": p.title,
                    "status": p.status,
                    "original_title": p.original_title,
                    "new_worldview": p.new_worldview,
                    "created_at": p.created_at.isoformat()
                }
                for p in projects
            ]
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
